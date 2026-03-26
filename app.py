#!/usr/bin/env python3
"""
Image Scraper UI - Web interface for Hermes B2B Image Scraper
=============================================================
Flask app with Server-Sent Events for real-time progress.

Usage:
    python app.py                    # Start on port 8787
    python app.py --port 9090        # Custom port
"""

import argparse
import base64
import hashlib
import io
import json
import logging
import os
import queue
import re
import sqlite3
import threading
import time
import uuid
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from flask import Flask, Response, jsonify, request, send_from_directory
from PIL import Image, ImageOps, ImageFilter

# ─── APP SETUP ────────────────────────────────────────────────────────────

app = Flask(__name__)
logger = logging.getLogger("scraper-ui")
logging.basicConfig(level=logging.INFO)

OUTPUT_DIR = Path("./output")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
)

# Active jobs: job_id -> {thread, queue, status, config}
active_jobs = {}


# ─── IMAGE QUALITY ANALYSIS ──────────────────────────────────────────────

class ImageQualityChecker:
    """Evaluates image quality on multiple dimensions (0-100 score)."""

    def __init__(self, config: dict):
        self.min_resolution = config.get("min_resolution", 200)
        self.min_quality_score = config.get("min_quality_score", 40)
        self.reject_watermarks = config.get("reject_watermarks", True)
        self.reject_blurry = config.get("reject_blurry", True)
        self.reject_multi_product = config.get("reject_multi_product", True)
        self.min_aspect_ratio = config.get("min_aspect_ratio", 0.4)
        self.max_aspect_ratio = config.get("max_aspect_ratio", 2.5)

    def evaluate(self, data: bytes) -> dict:
        """Returns {score: 0-100, passed: bool, reasons: [...]}."""
        result = {"score": 0, "passed": False, "reasons": [], "details": {}}

        try:
            img = Image.open(io.BytesIO(data))
            img_copy = Image.open(io.BytesIO(data))
        except Exception:
            result["reasons"].append("Cannot open image")
            return result

        scores = []

        # 1. Resolution check
        w, h = img.size
        result["details"]["resolution"] = f"{w}x{h}"
        if w < self.min_resolution or h < self.min_resolution:
            result["reasons"].append(f"Too small: {w}x{h} (min {self.min_resolution}px)")
            res_score = 0
        else:
            res_score = min(100, (min(w, h) / max(self.min_resolution * 2, 400)) * 100)
        scores.append(("resolution", res_score, 0.3))

        # 2. Aspect ratio check
        aspect = w / h if h > 0 else 0
        result["details"]["aspect_ratio"] = round(aspect, 2)
        if aspect < self.min_aspect_ratio or aspect > self.max_aspect_ratio:
            result["reasons"].append(f"Bad aspect ratio: {aspect:.2f}")
            aspect_score = 20
        else:
            # Prefer square-ish images for product photos
            ideal_deviation = abs(1.0 - aspect)
            aspect_score = max(20, 100 - ideal_deviation * 80)
        scores.append(("aspect_ratio", aspect_score, 0.15))

        # 3. Sharpness / blur detection (Laplacian variance)
        if self.reject_blurry:
            try:
                gray = img_copy.convert("L")
                laplacian = gray.filter(ImageFilter.FIND_EDGES)
                import numpy as np
                arr = np.array(laplacian, dtype=float)
                variance = arr.var()
                result["details"]["sharpness"] = round(variance, 1)
                if variance < 50:
                    result["reasons"].append(f"Blurry (sharpness: {variance:.0f})")
                    sharp_score = 10
                elif variance < 200:
                    sharp_score = 40 + (variance - 50) / 150 * 40
                else:
                    sharp_score = min(100, 80 + (variance - 200) / 500 * 20)
            except ImportError:
                sharp_score = 60  # Can't check without numpy
            scores.append(("sharpness", sharp_score, 0.25))

        # 4. Color diversity (simple check for placeholder/blank images)
        try:
            small = img_copy.resize((50, 50)).convert("RGB")
            colors = small.getcolors(maxcolors=2500) or []
            unique_colors = len(colors)
            result["details"]["unique_colors"] = unique_colors
            if unique_colors < 20:
                result["reasons"].append("Too uniform (possible placeholder)")
                color_score = 10
            elif unique_colors < 100:
                color_score = 30 + (unique_colors / 100) * 40
            else:
                color_score = min(100, 70 + (unique_colors / 500) * 30)
        except Exception:
            color_score = 50
        scores.append(("color_diversity", color_score, 0.15))

        # 5. File size adequacy
        file_size = len(data)
        result["details"]["file_size_kb"] = round(file_size / 1024, 1)
        if file_size < 5000:
            result["reasons"].append(f"File too small: {file_size / 1024:.1f}KB")
            size_score = 10
        elif file_size < 20000:
            size_score = 40
        else:
            size_score = min(100, 60 + (file_size / 200000) * 40)
        scores.append(("file_size", size_score, 0.15))

        # 6. Multi-product detection (heuristic, informational only)
        # The actual rejection is done by Claude Vision API later in the pipeline,
        # which is far more reliable. This is kept for logging/debugging.
        if self.reject_multi_product:
            try:
                multi_hint = self._detect_multi_product(img)
                result["details"]["multi_product_hint"] = multi_hint
            except Exception:
                pass

        # Weighted total
        total = sum(score * weight for _, score, weight in scores)
        result["score"] = round(total)
        result["passed"] = (total >= self.min_quality_score
                            and len([r for r in result["reasons"]
                                     if "Too small" in r or "Cannot" in r]) == 0)
        result["details"]["component_scores"] = {
            name: round(score, 1) for name, score, _ in scores
        }

        return result

    def _detect_multi_product(self, img: Image.Image) -> bool:
        """
        Detect if an image contains multiple separate product items.
        Uses multiple heuristics — any positive signal = reject.
        Returns True if multiple products are likely.
        """
        import numpy as np

        w, h = img.size

        # --- Heuristic 1: Wide aspect ratio ---
        # Product photos with multiple items are almost always wide/landscape.
        # A single product bag is typically portrait or squarish.
        aspect = w / h if h > 0 else 1.0
        if aspect > 1.8:
            # Very wide image — almost certainly multiple products or a banner
            return True

        # Work on a smaller copy for speed
        thumb = img.copy()
        thumb.thumbnail((300, 300), Image.LANCZOS)
        tw, th = thumb.size
        gray = np.array(thumb.convert("L"), dtype=np.float64)

        # --- Heuristic 2: Vertical column analysis ---
        # Split image into vertical columns. For each column, measure
        # edge density. Look for "gaps" (low-edge columns) that separate products.
        col_edges = np.abs(np.diff(gray, axis=1)).mean(axis=0)
        if len(col_edges) < 20:
            return False

        # Smooth the column edge profile
        k = max(3, tw // 25)
        if k % 2 == 0:
            k += 1
        kernel = np.ones(k) / k
        smoothed = np.convolve(col_edges, kernel, mode='valid')
        if len(smoothed) < 10:
            return False

        # Find how many "object peaks" there are separated by valleys
        mean_edge = smoothed.mean()
        if mean_edge < 3:
            return False  # Very uniform image, likely not products

        # Threshold: columns with edge density < 40% of mean are "gaps"
        gap_threshold = mean_edge * 0.35

        # Count transitions from "content" to "gap" — each gap = boundary between objects
        in_content = False
        content_regions = 0
        region_start = 0
        min_region_width = len(smoothed) * 0.1  # Region must be at least 10% of width

        for ci, val in enumerate(smoothed):
            if val > gap_threshold:
                if not in_content:
                    in_content = True
                    region_start = ci
            else:
                if in_content:
                    in_content = False
                    region_width = ci - region_start
                    if region_width >= min_region_width:
                        content_regions += 1

        # Count last region if we ended in content
        if in_content:
            region_width = len(smoothed) - region_start
            if region_width >= min_region_width:
                content_regions += 1

        if content_regions >= 3:
            return True

        # --- Heuristic 3: Repetitive pattern detection ---
        # If the image has visually repetitive vertical sections,
        # it's likely showing the same product multiple times or a lineup.
        if tw > 100 and aspect > 1.3:
            # Split into equal vertical slices and compare their histograms
            n_slices = min(4, max(2, int(aspect + 0.5)))
            slice_w = tw // n_slices
            if slice_w > 20:
                histograms = []
                for i in range(n_slices):
                    slc = gray[:, i * slice_w:(i + 1) * slice_w]
                    hist, _ = np.histogram(slc.flatten(), bins=32, range=(0, 256))
                    hist = hist.astype(float)
                    norm = hist.sum()
                    if norm > 0:
                        hist /= norm
                    histograms.append(hist)

                # Compare adjacent slices — if they're all similar, it's repetitive
                similarities = []
                for i in range(len(histograms) - 1):
                    # Bhattacharyya coefficient (1 = identical, 0 = different)
                    bc = np.sum(np.sqrt(histograms[i] * histograms[i + 1]))
                    similarities.append(bc)

                if len(similarities) >= 2 and all(s > 0.92 for s in similarities):
                    # All slices look very similar + wide image = likely product lineup
                    if aspect > 1.4:
                        return True

        return False


# ─── RELEVANCE CHECKER (CLIP) ─────────────────────────────────────────────

class RelevanceChecker:
    """
    Uses OpenCLIP to compare image content with product description.
    Lazy-loads the model on first use to avoid slowing down startup.
    Returns a similarity score 0-100.
    """

    def __init__(self, min_relevance: float = 0.20):
        self.min_relevance = min_relevance  # cosine similarity threshold
        self.model = None
        self.preprocess = None
        self.tokenizer = None
        self._loaded = False
        self._available = None

    def _load(self):
        if self._loaded:
            return
        try:
            import open_clip
            import torch
            self.torch = torch  # Store reference for use in check()

            # Use a small, fast model - ViT-B-32 is a good balance
            self.model, _, self.preprocess = open_clip.create_model_and_transforms(
                'ViT-B-32', pretrained='laion2b_s34b_b79k'
            )
            self.tokenizer = open_clip.get_tokenizer('ViT-B-32')
            self.model.eval()

            # Self-test: run a dummy inference to catch runtime errors early
            dummy_img = Image.new("RGB", (224, 224), (128, 128, 128))
            dummy_input = self.preprocess(dummy_img).unsqueeze(0)
            dummy_text = self.tokenizer(["test product"])
            with torch.no_grad():
                self.model.encode_image(dummy_input)
                self.model.encode_text(dummy_text)

            self._loaded = True
            self._available = True
            logger.info("CLIP model loaded and self-test passed")
        except Exception as e:
            logger.warning(f"CLIP not available: {type(e).__name__}: {e}")
            self._loaded = True
            self._available = False

    @property
    def available(self) -> bool:
        if self._available is None:
            self._load()
        return self._available

    def _build_clip_prompts(self, product_name: str) -> list[str]:
        """
        Build multiple CLIP-friendly text prompts from a product name.
        CLIP works best with descriptive phrases like "a photo of X".
        We generate several variants to maximize matching.
        """
        cleaned = clean_product_query(product_name)

        # Remove generic RO words, keep meaningful keywords
        skip = {"ceai", "cafea", "de", "cu", "si", "din", "la", "pt", "pentru",
                "infuzie", "buc", "cutie", "pachet", "set"}
        key_words = [w for w in cleaned.split() if w.lower() not in skip and len(w) > 1]
        key_phrase = " ".join(key_words) if key_words else cleaned

        # Translate RO product type to EN for CLIP (trained mostly on English)
        en_phrase = cleaned
        for ro, en in _RO_EN_MAP.items():
            en_phrase = re.sub(r'\b' + ro + r'\b', en, en_phrase, flags=re.IGNORECASE)

        prompts = [
            f"a product photo of {key_phrase}",
            f"a photo of {en_phrase}",
            f"{key_phrase} product packaging",
            cleaned,
        ]
        # Dedupe
        return list(dict.fromkeys(prompts))

    def check(self, image_data: bytes, product_name: str) -> dict:
        """
        Compare image with product name using CLIP.
        Uses multiple text prompts and takes the best match.
        Returns {score: 0-100, similarity: float, relevant: bool}
        """
        if not self.available:
            return {"score": 50, "similarity": 0.0, "relevant": True, "reason": "CLIP not available"}

        try:
            torch = self.torch

            # Prepare image
            img = Image.open(io.BytesIO(image_data)).convert("RGB")
            image_input = self.preprocess(img).unsqueeze(0)

            # Build multiple text prompts for better matching
            prompts = self._build_clip_prompts(product_name)

            # Tokenize all prompts at once
            text_input = self.tokenizer(prompts)

            # Compute similarity with all prompts, take the best
            with torch.no_grad():
                image_features = self.model.encode_image(image_input)
                text_features = self.model.encode_text(text_input)

                image_features = image_features / image_features.norm(dim=-1, keepdim=True)
                text_features = text_features / text_features.norm(dim=-1, keepdim=True)

                similarities = (image_features @ text_features.T).squeeze()

                # Take the maximum similarity across all prompts
                if similarities.dim() == 0:
                    similarity = similarities.item()
                else:
                    similarity = similarities.max().item()

            score = max(0, min(100, (similarity - 0.10) / 0.30 * 100))
            relevant = similarity >= self.min_relevance

            reason = ""
            if not relevant:
                reason = f"Low relevance ({similarity:.2f} < {self.min_relevance})"

            # CLIP-based multi-product check:
            # Compare "single product" vs "multiple products" prompts
            multi_product = False
            try:
                single_prompts = self.tokenizer([
                    "a single product on white background",
                    "one item isolated on plain background",
                    "a photo of one product package",
                    "a single bag of coffee",
                ])
                multi_prompts = self.tokenizer([
                    "multiple products side by side",
                    "several bags of coffee in a row",
                    "a group of products displayed together",
                    "many packages arranged together on a shelf",
                    "a collection of different product varieties",
                    "product lineup showing multiple items",
                ])
                with torch.no_grad():
                    single_feat = self.model.encode_text(single_prompts)
                    multi_feat = self.model.encode_text(multi_prompts)
                    single_feat = single_feat / single_feat.norm(dim=-1, keepdim=True)
                    multi_feat = multi_feat / multi_feat.norm(dim=-1, keepdim=True)

                    single_sim = (image_features @ single_feat.T).squeeze().max().item()
                    multi_sim = (image_features @ multi_feat.T).squeeze().max().item()

                logger.info(f"CLIP multi-product check: single={single_sim:.3f} multi={multi_sim:.3f}")
                # Flag as multi-product if:
                # 1. multi > single (even slightly), OR
                # 2. multi score is high on its own (> 0.26)
                if multi_sim > single_sim or multi_sim > 0.26:
                    multi_product = True
                    logger.info(f"CLIP multi-product DETECTED: single={single_sim:.3f} multi={multi_sim:.3f}")
            except Exception as e:
                logger.debug(f"CLIP multi-product check failed: {e}")

            return {
                "score": round(score),
                "similarity": round(similarity, 3),
                "relevant": relevant,
                "reason": reason,
                "multi_product": multi_product,
            }
        except Exception as e:
            logger.warning(f"CLIP check error for '{product_name[:30]}': {type(e).__name__}: {e}")
            import traceback
            logger.warning(traceback.format_exc())
            return {"score": 50, "similarity": 0.0, "relevant": True, "reason": f"Error: {e}"}


# Global instance (lazy loaded)
_relevance_checker = None

def get_relevance_checker(min_relevance: float = 0.20) -> RelevanceChecker:
    global _relevance_checker
    if _relevance_checker is None:
        _relevance_checker = RelevanceChecker(min_relevance)
    return _relevance_checker


# ─── CLAUDE VISION: SINGLE PRODUCT CHECK ─────────────────────────────────

_gemini_last_call = 0.0  # timestamp of last Gemini API call

def gemini_vision_check_image(image_data: bytes, product_name: str, api_key: str) -> dict:
    """
    Use Google Gemini Flash Vision (FREE tier) to verify a product image:
    1. Is it exactly ONE product? (not multiple items)
    2. Does the text on the packaging match the product name?

    Gemini Flash free tier: 15 RPM, 1M tokens/day.
    Rate limited to ~1 call per 4.5 seconds to stay within limits.

    Returns {
        "ok": bool,           # True = image is good to use
        "count": int,         # Number of items detected
        "text_match": bool,   # True = packaging text matches product name
        "visible_text": str,  # Text read from the packaging
        "reason": str,        # Human-readable explanation
    }
    """
    if not api_key:
        return {"ok": True, "count": 1, "text_match": True, "visible_text": "",
                "reason": "No Gemini API key, skipping check"}

    # Rate limiting: flash-lite has 30 RPM free tier → 1 call per 2s
    global _gemini_last_call
    elapsed = time.time() - _gemini_last_call
    if elapsed < 2.0:
        time.sleep(2.0 - elapsed)
    _gemini_last_call = time.time()

    try:
        # Log image size for debugging
        try:
            from PIL import Image
            import io
            img_pil = Image.open(io.BytesIO(image_data))
            w, h = img_pil.size
            logger.info(f"Gemini Vision input image: {w}x{h} ({len(image_data)} bytes)")

            # If image is very small, upscale to help Gemini read text
            if w < 400 or h < 400:
                scale = max(400 / w, 400 / h)
                new_w, new_h = int(w * scale), int(h * scale)
                img_pil = img_pil.resize((new_w, new_h), Image.LANCZOS)
                buf = io.BytesIO()
                fmt = "PNG" if image_data[:4] == b'\x89PNG' else "JPEG"
                img_pil.save(buf, format=fmt, quality=95)
                image_data = buf.getvalue()
                logger.info(f"Gemini Vision upscaled to {new_w}x{new_h}")
        except Exception as e:
            logger.debug(f"Could not check image dimensions: {e}")

        # Convert image to base64
        img_b64 = base64.b64encode(image_data).decode("utf-8")

        # Detect media type
        media_type = "image/jpeg"
        if image_data[:4] == b'\x89PNG':
            media_type = "image/png"
        elif image_data[:4] == b'RIFF':
            media_type = "image/webp"

        # Try models in order — fallback if quota exhausted (429)
        _gemini_models = ["gemini-2.5-flash-lite", "gemini-2.5-flash", "gemini-2.0-flash"]

        payload = {
            "contents": [{
                "parts": [
                    {
                        "inline_data": {
                            "mime_type": media_type,
                            "data": img_b64,
                        },
                    },
                    {
                        "text": (
                            "Answer these two questions about this product image.\n\n"
                            "1) How many individual physical packages (bags, boxes, bottles, cans) "
                            "are visible? Count each bag separately even if they are the same product. "
                            "If you see 4 bags next to each other, answer 4.\n\n"
                            "2) Read ALL text visible on the product packaging and tell me the "
                            "brand name and product name exactly as written.\n\n"
                            "Format your answer EXACTLY like this:\n"
                            "COUNT: 1\n"
                            "TEXT: Lavazza Crema e Aroma"
                        ),
                    },
                ],
            }],
            "generationConfig": {
                "maxOutputTokens": 400,
                "temperature": 0.1,
            },
        }

        resp = None
        for model_name in _gemini_models:
            api_url = (
                f"https://generativelanguage.googleapis.com/v1beta/"
                f"models/{model_name}:generateContent"
                f"?key={api_key}"
            )
            resp = requests.post(
                api_url,
                headers={"Content-Type": "application/json"},
                json=payload,
                timeout=30,
            )
            if resp.status_code == 429:
                logger.info(f"Gemini Vision: {model_name} quota exhausted, trying next model...")
                continue
            break

        resp.raise_for_status()
        resp_data = resp.json()

        # Extract answer from Gemini response
        # Gemini 2.5 Flash may return multiple parts (thinking + answer)
        answer = ""
        try:
            parts = resp_data["candidates"][0]["content"]["parts"]
            # Concatenate all text parts (skip thinking parts if any)
            texts = []
            for p in parts:
                if "text" in p:
                    texts.append(p["text"].strip())
            answer = "\n".join(texts)
            logger.info(f"Gemini Vision got {len(parts)} parts, combined len={len(answer)}")
        except (KeyError, IndexError):
            logger.warning(f"Gemini Vision unexpected response: {json.dumps(resp_data)[:300]}")
            return {"ok": True, "count": -1, "text_match": True, "visible_text": "",
                    "reason": "Unexpected response format"}

        logger.info(f"Gemini Vision raw answer: {answer}")

        # Parse COUNT
        count = 1
        count_m = re.search(r'COUNT:\s*(\d+)', answer, re.IGNORECASE)
        if count_m:
            count = int(count_m.group(1))

        # Parse TEXT
        visible_text = ""
        text_m = re.search(r'TEXT:\s*(.*)', answer, re.IGNORECASE)
        if text_m:
            visible_text = text_m.group(1).strip().split('\n')[0]  # First line only
        logger.info(f"Gemini Vision parsed: count={count}, text='{visible_text[:80]}'")


        # Check if visible text matches product name
        clean_name = re.sub(r'\b\d+\s*(kg|g|ml|l|cl|gr)\b', '', product_name, flags=re.IGNORECASE)
        clean_name = re.sub(r'\s+', ' ', clean_name).strip()
        name_words = set(re.findall(r'[a-zA-Z]{3,}', clean_name.lower()))
        _skip_vision = {"cafea", "coffee", "ceai", "tea", "boabe", "beans",
                        "capsule", "capsules", "negru", "albastru", "expert",
                        "alb", "rosu"}
        distinctive_words = {w for w in name_words if w not in _skip_vision}

        visible_lower = visible_text.lower()
        matching_words = {w for w in distinctive_words if w in visible_lower}

        # Need enough distinctive words to match:
        # - Brand words (e.g. "lavazza") alone are NOT enough
        # - Need at least the brand + one product-specific word
        # - e.g. for "LAVAZZA Tierra India": need "lavazza" + at least one of {"tierra","india"}
        n_dist = len(distinctive_words)
        if n_dist <= 1:
            min_required = n_dist  # If only 1 word, need it
        elif n_dist == 2:
            min_required = 2  # Need both
        else:
            min_required = max(2, (n_dist + 1) // 2)  # Need majority, at least 2

        text_match = len(matching_words) >= min_required if distinctive_words else True

        # If Gemini returned empty text, we can't verify — mark as uncertain
        text_uncertain = not visible_text.strip()
        if text_uncertain:
            text_match = True  # Don't reject based on empty text alone

        is_single = count == 1
        ok = is_single and text_match

        reason_parts = []
        if not is_single:
            reason_parts.append(f"{count} items (need 1)")
        if not text_match:
            missing = distinctive_words - matching_words
            reason_parts.append(f"text mismatch: saw '{visible_text[:60]}', missing {missing}")
        reason = "; ".join(reason_parts) if reason_parts else "OK"

        logger.info(f"Gemini Vision: count={count}, text='{visible_text[:50]}', "
                     f"match={matching_words}/{distinctive_words} → ok={ok}")

        return {
            "ok": ok,
            "count": count,
            "text_match": text_match,
            "visible_text": visible_text,
            "matching_words": matching_words,
            "reason": reason,
        }

    except requests.exceptions.HTTPError as e:
        err_text = ""
        try:
            err_text = e.response.text[:200]
        except Exception:
            pass
        logger.warning(f"Gemini Vision API error: {e} - {err_text}")
        # Return "fallback" — let pipeline use local checks (CLIP + conflict) instead
        return {"ok": None, "count": -1, "text_match": None, "visible_text": "",
                "reason": f"API error, using local fallback"}
    except Exception as e:
        logger.debug(f"Gemini Vision check failed: {e}")
        return {"ok": None, "count": -1, "text_match": None, "visible_text": "",
                "reason": f"Vision failed, using local fallback"}


# ─── SEARCH ENGINES ───────────────────────────────────────────────────────

class BingImageScraper:
    """Scrape Bing Images - no API key needed."""

    SEARCH_URL = "https://www.bing.com/images/search"

    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": USER_AGENT,
            "Accept": "text/html,application/xhtml+xml",
            "Accept-Language": "en-US,en;q=0.9,ro;q=0.8",
        })
        self._last_call = 0

    def search(self, query: str, max_results: int = 8) -> list[dict]:
        elapsed = time.time() - self._last_call
        if elapsed < 0.5:
            time.sleep(0.5 - elapsed)

        try:
            resp = self.session.get(
                self.SEARCH_URL,
                params={"q": query, "form": "HDRSC2", "first": "1",
                        "qft": "+filterui:imagesize-large"},
                timeout=15,
            )
            self._last_call = time.time()
            resp.raise_for_status()
        except Exception as e:
            logger.debug(f"Bing scrape error: {e}")
            return []

        results = []
        urls = re.findall(r'murl&quot;:&quot;(https?://[^&]+?)&quot;', resp.text)

        for url in urls[:max_results * 2]:
            if len(results) >= max_results:
                break
            if any(skip in url.lower() for skip in [
                'favicon', 'logo', '1x1', 'pixel', '.svg', '.gif', 'placeholder'
            ]):
                continue
            results.append({"image": url, "title": query, "source": "bing"})

        return results


class DuckDuckGoSearch:
    """DuckDuckGo image search - free, no API key."""

    def __init__(self):
        self._last_call = 0

    def search(self, query: str, max_results: int = 8) -> list[dict]:
        elapsed = time.time() - self._last_call
        if elapsed < 4:
            time.sleep(4 - elapsed)

        try:
            try:
                from ddgs import DDGS
            except ImportError:
                from duckduckgo_search import DDGS
            with DDGS() as ddgs:
                results = list(ddgs.images(
                    keywords=query, region="wt-wt", safesearch="moderate",
                    size="Large", type_image="photo", max_results=max_results,
                ))
            self._last_call = time.time()
            for r in results:
                r["source"] = "ddg"
            return results
        except Exception as e:
            logger.debug(f"DDG error: {e}")
            return []


class PexelsSearch:
    """Pexels API - 200 requests/hour, free."""

    BASE_URL = "https://api.pexels.com/v1/search"

    def __init__(self, api_key: str):
        self.api_key = api_key
        self.session = requests.Session()
        self.session.headers["Authorization"] = api_key

    def search(self, query: str, max_results: int = 8) -> list[dict]:
        if not self.api_key:
            return []
        try:
            resp = self.session.get(
                self.BASE_URL,
                params={"query": query, "per_page": min(max_results, 15),
                        "orientation": "square"},
                timeout=15,
            )
            if resp.status_code == 429:
                time.sleep(60)
                return self.search(query, max_results)
            resp.raise_for_status()
            data = resp.json()
            results = []
            for photo in data.get("photos", []):
                src = photo.get("src", {})
                img_url = src.get("medium") or src.get("large") or src.get("original", "")
                if img_url:
                    results.append({
                        "image": img_url, "title": photo.get("alt", ""),
                        "source": "pexels",
                    })
            return results
        except Exception as e:
            logger.debug(f"Pexels error: {e}")
            return []

    @property
    def available(self):
        return bool(self.api_key)


class AIProductMatcher:
    """
    AI-powered product matching: given a product name and a list of candidates
    from a site, pick the best match. Uses Anthropic Claude API if available,
    otherwise falls back to an advanced local matching algorithm.
    """

    def __init__(self, api_key: str = ""):
        self.api_key = api_key or os.environ.get("ANTHROPIC_API_KEY", "")
        self._available = None

    @property
    def available(self):
        if self._available is None:
            self._available = bool(self.api_key)
        return self._available

    def match(self, product_name: str, candidates: list[dict],
              max_results: int = 3) -> list[dict]:
        """
        Match product_name against candidates (each having 'title' and 'url').
        Returns the best matching candidates, sorted by relevance.

        Args:
            product_name: The product we're looking for (e.g. "Pireu MONIN Banane 1L")
            candidates: List of dicts with at least 'title' and 'url' keys
            max_results: How many top matches to return

        Returns:
            Sorted list of best matching candidates
        """
        if not candidates:
            return []

        # Detect product type from the query (e.g. "pireu" → "puree")
        # Used as a HARD FILTER on all matcher results (AI and local)
        name_type = self._detect_product_type(product_name)
        print(f"[MATCHER] product='{product_name}' | type={name_type} | candidates={len(candidates)} | ai_available={self.available}")
        for i, c in enumerate(candidates[:8]):
            ct = self._detect_product_type(f"{c.get('title','')} {c.get('url','')}")
            print(f"  [{i}] type={ct} | {c.get('title','(no title)')} | {c.get('url','')[:80]}")

        # Try AI matching first if API key available
        if self.available:
            try:
                ai_result = self._ai_match(product_name, candidates, max_results)
                if ai_result:
                    # POST-FILTER: remove type-mismatched results from AI output
                    # e.g. if searching for "pireu" (puree), reject any "sirop" (syrup) results
                    if name_type:
                        filtered = []
                        for r in ai_result:
                            cand_text = f"{r.get('title','')} {r.get('url','')}"
                            cand_type = self._detect_product_type(cand_text)
                            if cand_type and cand_type != name_type:
                                logger.info(f"Type filter rejected AI result: '{r.get('title','')}' (wanted={name_type}, got={cand_type})")
                            else:
                                filtered.append(r)
                        ai_result = filtered
                    if ai_result:
                        logger.info(f"AI API matched '{product_name}' → '{ai_result[0].get('title', ai_result[0].get('url', ''))}'")
                        return ai_result
                    else:
                        logger.info(f"AI API results rejected by type filter for '{product_name}' (wanted type={name_type}), trying local")
                else:
                    logger.info(f"AI API returned no match for '{product_name}', trying local matcher")
            except Exception as e:
                logger.warning(f"AI API matching failed for '{product_name}': {e}, falling back to local")

        # Fall back to advanced local matching (always available, has its own type filter)
        local_result = self._local_match(product_name, candidates, max_results)
        if local_result:
            logger.info(f"Local matched '{product_name}' → '{local_result[0].get('title', local_result[0].get('url', ''))}'")
        return local_result

    def _ai_match(self, product_name: str, candidates: list[dict],
                  max_results: int) -> list[dict]:
        """Use Anthropic Claude to intelligently match products."""
        # Build candidate list for the prompt
        candidate_lines = []
        for i, c in enumerate(candidates[:30]):  # Limit to 30 candidates
            title = c.get("title", c.get("url", ""))
            candidate_lines.append(f"{i}. {title}")

        prompt = f"""You are a product matching system. Given a product name and a list of candidates from an e-commerce site, identify which candidates are the SAME product.

Product to find: "{product_name}"

Candidates:
{chr(10).join(candidate_lines)}

Rules:
- Match by brand, flavor/variant, and product type
- "Pireu"/"Piure"/"Puree" are the same thing (puree)
- "Banane"/"Banana" are the same (banana)
- "Sirop"/"Syrup"/"Sirup" are the same (syrup)
- CRITICAL: The product TYPE must match. If searching for "Pireu" (puree), do NOT return "Sirop" (syrup) products, and vice versa. Type mismatch = no match.
- Ignore packaging sizes (1L, 0.7L, 250g)
- A product must match on BRAND + FLAVOR + TYPE to be correct
- Return ONLY the index numbers of matching candidates, comma-separated
- If no match found, return "NONE"

Best matches (indices only):"""

        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": self.api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
            json={
                "model": "claude-haiku-4-5-20251001",
                "max_tokens": 100,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=10,
        )
        resp.raise_for_status()
        answer = resp.json()["content"][0]["text"].strip()

        if answer.upper() == "NONE":
            return []

        # Parse indices
        indices = []
        for part in answer.replace(",", " ").split():
            try:
                idx = int(part.strip("."))
                if 0 <= idx < len(candidates):
                    indices.append(idx)
            except ValueError:
                continue

        return [candidates[i] for i in indices[:max_results]]

    # Product type words — used to detect type mismatches (pireu vs sirop)
    _PRODUCT_TYPES = {
        "puree":    {"piure", "pireu", "puree", "purée", "püree", "puré"},
        "syrup":    {"sirop", "syrup", "sirup", "sirope", "sciroppo"},
        "tea":      {"ceai", "tea", "tee", "té", "tè", "thé"},
        "coffee":   {"cafea", "coffee", "kaffee", "café"},
        "juice":    {"suc", "juice", "saft", "jus", "jugo", "succo"},
        "sauce":    {"sos", "sauce", "soße"},
        "liqueur":  {"lichior", "liqueur", "likör", "licor"},
        "cream":    {"crema", "cream", "crème", "frisca"},
        "infusion": {"infuzie", "infusion", "aufguss"},
        "concentrate": {"concentrat", "concentrate"},
    }

    # Words that when preceding "cream"/"crema" indicate it's a FLAVOR name,
    # not a product type. "Blueberry Cream" = flavor, "Cream" alone = product type.
    _CREAM_FLAVOR_PREFIXES = {
        "blueberry", "strawberry", "raspberry", "vanilla", "chocolate", "caramel",
        "banana", "mango", "coconut", "hazelnut", "pistachio", "cookie", "cookies",
        "ice", "irish", "butterscotch", "toffee", "peach", "lemon", "orange",
        "cherry", "apple", "peanut", "almond", "mint", "coffee", "mocha",
        "afine", "capsuni", "zmeura", "vanilie", "ciocolata", "banane",
    }

    # Known coffee brand product line names containing "crema"/"cream"
    # These are product names, NOT product types
    _CREAM_PRODUCT_NAMES = {
        "super crema", "crema e aroma", "crema & aroma", "crema e gusto",
        "crema & gusto", "pronto crema", "caffe crema", "caffè crema",
        "crema ricca", "crema classica", "crema dolce", "gran crema",
        "crema gustoso", "crema leggero", "crema intenso", "espresso cremoso",
        "crema aroma", "crema gusto",
    }

    def _detect_product_type(self, text: str) -> str | None:
        """Detect product type category from text. Returns type key or None.
        Context-aware: 'Blueberry Cream' is a flavor, not product type 'cream'.
        Also: 'Super Crema' / 'Crema e Aroma' are coffee product names, not cream.
        """
        text_lower = text.lower()
        tokens = set(re.findall(r'[a-zăâîșțéèêëüöïôàáùúñçÀ-ÿ]+', text_lower))

        # Collect all matching types
        matched_types = {}
        for type_key, type_words in self._PRODUCT_TYPES.items():
            matched = tokens & type_words
            if matched:
                matched_types[type_key] = matched

        if not matched_types:
            return None

        # Special handling for "cream"/"crema": skip if it's a known product name
        # like "Super Crema", "Crema e Aroma", etc.
        if "cream" in matched_types:
            # Check if it's a known coffee product line name
            for pname in self._CREAM_PRODUCT_NAMES:
                if pname in text_lower:
                    del matched_types["cream"]
                    break

            # If "cream" is still there, check flavor prefix rule
            if "cream" in matched_types:
                is_flavor_cream = False
                for cream_word in matched_types["cream"]:
                    pattern = r'(\w+)\s+' + re.escape(cream_word)
                    for m in re.finditer(pattern, text_lower):
                        preceding = m.group(1)
                        if preceding in self._CREAM_FLAVOR_PREFIXES:
                            is_flavor_cream = True
                            break
                    if is_flavor_cream:
                        break
                if is_flavor_cream:
                    del matched_types["cream"]

            # If both "cream" and "coffee" match, prefer coffee
            # (e.g. "cafea lavazza crema" = coffee, not cream)
            if "cream" in matched_types and "coffee" in matched_types:
                del matched_types["cream"]

        if not matched_types:
            return None

        # Return first remaining match (priority order from dict)
        return next(iter(matched_types))


    def _local_match(self, product_name: str, candidates: list[dict],
                     max_results: int) -> list[dict]:
        """
        Advanced local matching using synonym-aware word overlap + fuzzy scoring.
        Much smarter than simple keyword matching.
        """
        # Tokenize and normalize the product name
        name_tokens = self._tokenize(product_name)
        # Separate into brand tokens and product tokens
        name_brand = []
        name_product = []
        for tok in name_tokens:
            en = normalize_to_english(tok)
            if en and en != tok:
                name_product.append(tok)
            else:
                name_brand.append(tok)

        # Detect the product type from the input name (e.g. "pireu" → "puree")
        name_type = self._detect_product_type(product_name)

        scored = []
        for c in candidates:
            title = c.get("title", "")
            url_path = urlparse(c.get("url", "")).path
            url_text = url_path.replace("-", " ").replace("_", " ").replace("/", " ")
            cand_text = f"{title} {url_text}".lower()
            cand_tokens = set(self._tokenize(cand_text))

            score = 0.0
            brand_matched = False

            # Brand matching (exact match required, high weight)
            for bt in name_brand:
                if bt in cand_tokens:
                    score += 10.0
                    brand_matched = True
                elif any(bt in ct for ct in cand_tokens):
                    score += 5.0
                    brand_matched = True

            # CRITICAL: if product has brand words and candidate doesn't match ANY,
            # skip this candidate entirely — it's a different brand's product
            if name_brand and not brand_matched:
                continue

            # Product word matching (synonym-aware, medium weight)
            for pt in name_product:
                variants = get_word_variants(pt)
                if variants & cand_tokens:
                    score += 6.0  # Exact or synonym match
                elif any(any(v in ct for ct in cand_tokens) for v in variants):
                    score += 3.0  # Partial variant match

            # TYPE MISMATCH: if product is "pireu" but candidate says "sirop", SKIP entirely.
            # This is as critical as brand mismatch — wrong product type = wrong product.
            # e.g. "Monin Green Apple Sirop 0.7L" is NOT "Pireu MONIN Green Apple 1L"
            if name_type:
                cand_type = self._detect_product_type(cand_text)
                if cand_type and cand_type != name_type:
                    print(f"[LOCAL] SKIP type mismatch: '{title}' (wanted={name_type}, got={cand_type})")
                    continue  # Hard skip — wrong product type entirely
                elif cand_type == name_type:
                    print(f"[LOCAL] TYPE OK: '{title}' (type={cand_type}, score so far={score})")

            # Penalty: candidate has extra brand-like words (different product line)
            cand_brand_tokens = [ct for ct in cand_tokens
                                 if not normalize_to_english(ct) or normalize_to_english(ct) == ct]
            noise_words = {"de", "cu", "si", "din", "la", "pt", "pentru", "fructe", "fruit",
                           "piure", "puree", "sirop", "syrup", "ceai", "tea", "infuzie",
                           "looseleaf", "herbal", "premium", "original"}
            extra_brands = set(cand_brand_tokens) - set(name_brand) - noise_words
            if extra_brands and name_brand:
                score -= len(extra_brands) * 1.0

            # WEIGHT/VOLUME matching bonus: prefer candidates with matching weight
            # Extract weight from product name and candidate
            name_vol = re.search(r'(\d+(?:[.,]\d+)?)\s*(l|ml|gr?|kg|cl)\b', product_name, re.IGNORECASE)
            cand_vol = re.search(r'(\d+(?:[.,]\d+)?)\s*(l|ml|gr?|kg|cl)\b', cand_text, re.IGNORECASE)
            if name_vol and cand_vol:
                # Normalize volumes to common unit (ml for liquids, g for weight)
                def _to_base(num_str, unit):
                    n = float(num_str.replace(",", "."))
                    u = unit.lower().rstrip("r")  # "gr" → "g"
                    if u == "l": return n * 1000, "ml"
                    if u == "cl": return n * 10, "ml"
                    if u == "ml": return n, "ml"
                    if u == "kg": return n * 1000, "g"
                    if u == "g": return n, "g"
                    return n, u

                n_val, n_unit = _to_base(name_vol.group(1), name_vol.group(2))
                c_val, c_unit = _to_base(cand_vol.group(1), cand_vol.group(2))
                if n_unit == c_unit:
                    if abs(n_val - c_val) < 0.1:
                        score += 8.0  # Exact weight match — strong signal
                    else:
                        score -= 2.0  # Different weight — mild penalty

            # INFUZIE ↔ LOOSELEAF equivalence bonus
            _product_form_groups = [
                {"infuzie", "infusion", "looseleaf", "loose", "frunze", "leaves", "leaf"},
                {"plic", "pliculete", "tea bag", "teabag", "sachet"},
            ]
            name_lower = product_name.lower()
            for group in _product_form_groups:
                name_has = any(w in name_lower for w in group)
                cand_has = any(w in cand_text for w in group)
                if name_has and cand_has:
                    score += 5.0  # Form match (both looseleaf/infuzie)
                elif name_has and not cand_has:
                    score -= 3.0  # Name says infuzie but candidate doesn't have it

            if score > 0:
                scored.append((score, c))

        scored.sort(key=lambda x: x[0], reverse=True)
        return [c for _, c in scored[:max_results]]

    @staticmethod
    def _tokenize(text: str) -> list[str]:
        """Tokenize text into lowercase words, filtering noise."""
        words = re.findall(r'[a-zA-ZăâîșțéèêëüöïôàáùúñçÀ-ÿ]+', text.lower())
        noise = {"de", "cu", "si", "din", "la", "pt", "pentru", "the", "and", "for",
                 "per", "con", "del", "les", "des", "von", "und",
                 "mic", "mare", "mediu", "mini"}
        return [w for w in words if len(w) >= 2 and w not in noise]


class DirectSiteScraper:
    """
    Scrape product images directly from e-commerce sites.
    Instead of relying on Google/Bing to index the site, we:
    1. Hit the site's own search (common patterns: /catalogsearch/result/?q=, /search?q=, etc.)
    2. Parse the search results to find product pages
    3. Extract product images from those pages
    """

    # Common e-commerce search URL patterns
    # Searchanise API keys cache: domain → api_key (or None if not found)
    _searchanise_keys: dict[str, str | None] = {}
    # Brand page cache: domain → {brand_pattern_url: (entries_list, effective_base)}
    _brand_page_cache: dict[str, dict] = {}

    SEARCH_PATTERNS = [
        "/catalogsearch/result/?q={query}",    # Magento
        "/search?q={query}",                   # Shopify, generic
        "/search?s={query}",                   # WooCommerce
        "/cautare?q={query}",                  # Romanian sites
        "/?s={query}",                         # WordPress
        "/product/search?keyword={query}",     # Custom platforms (e.g. finestore.ro)
        "/products/search?q={query}",          # Some Shopify variants
        "/?dispatch=products.search&q={query}",  # CS-Cart basic
        # CS-Cart extended — many CS-Cart sites need subcats/pname/pkeywords flags
        "/?subcats=Y&pcode_from_q=Y&pshort=Y&pfull=Y&pname=Y&pkeywords=Y&search_performed=Y&cid=0&q={query}&dispatch=products.search",
    ]

    def __init__(self, ai_matcher: AIProductMatcher | None = None):
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": USER_AGENT,
            "Accept": "text/html,application/xhtml+xml",
            "Accept-Language": "en-US,en;q=0.9,ro;q=0.8",
        })
        self.ai_matcher = ai_matcher
        # Per-instance caches (reset per scraper instance)
        self._brand_page_cache = {}

    def _try_searchanise_api(self, site: str, query: str, product_name: str,
                             max_results: int = 5) -> list[dict]:
        """
        Try Searchanise cloud JSON API. Many e-commerce sites use Searchanise
        which renders results via JS. The cloud API returns structured JSON.
        Returns list of {"image": url, "title": ..., "source": "direct:site", "product_url": ...}
        """
        results = []
        base_url = f"https://{site}" if not site.startswith("http") else site

        # Get or discover Searchanise API key and host for this site
        if site not in self._searchanise_keys:
            try:
                resp = self.session.get(base_url, timeout=6)
                # Look for Searchanise config: {"host":"https://searchserverapi.com","api_key":"..."}
                import re as _re
                # Pattern: "api_key":"ALPHANUMERIC" near "Searchanise" or "searchserverapi"
                m = _re.search(
                    r'(?:Searchanise|searchserverapi)[^}]*?"api_key"\s*:\s*"([a-zA-Z0-9]{8,})"',
                    resp.text, _re.IGNORECASE
                )
                if not m:
                    m = _re.search(
                        r'"api_key"\s*:\s*"([a-zA-Z0-9]{8,})"[^}]*?(?:Searchanise|searchserverapi)',
                        resp.text, _re.IGNORECASE
                    )
                # Also find the API host
                host_m = _re.search(
                    r'(?:Searchanise|searchserverapi)[^}]*?"host"\s*:\s*"([^"]+)"',
                    resp.text, _re.IGNORECASE
                )
                api_key = m.group(1) if m else None
                api_host = host_m.group(1).replace("\\/", "/") if host_m else "https://www.searchanise.com"
                self._searchanise_keys[site] = (api_key, api_host) if api_key else (None, None)
                if api_key:
                    print(f"[SEARCHANISE] Found API key for {site}: {api_key} (host: {api_host})")
            except Exception:
                self._searchanise_keys[site] = (None, None)

        api_key, api_host = self._searchanise_keys.get(site, (None, None))
        if not api_key:
            return []

        # Query the Searchanise cloud API
        # Use the full product name for best results
        search_q = product_name or query
        api_url = f"{api_host}/getresults?api_key={api_key}&q={requests.utils.quote(search_q)}&maxResults=10"
        try:
            resp = self.session.get(api_url, timeout=6)
            if resp.status_code != 200:
                return []

            data = json.loads(resp.text)
            items = data.get("items", [])
            if not items:
                return []

            print(f"[SEARCHANISE] {site} query='{search_q[:50]}' → {len(items)} results")

            # Convert to entries for matcher
            entries = []
            for item in items:
                link = item.get("link", "")
                title = item.get("title", "")
                image = item.get("image_link", "")
                if link:
                    entries.append({
                        "url": link, "title": title, "image": image,
                        "alt": title, "all_text": title,
                    })

            if not entries:
                return []

            # Pre-filter with distinctive words — require ALL distinctive words
            if product_name:
                pn_words = set(re.findall(r'[a-z]{4,}', product_name.lower()))
                _skip_sa = set(_RO_EN_MAP.keys()) | set(_RO_EN_MAP.values()) | {
                    "cafea", "coffee", "ceai", "tea", "boabe", "beans",
                    "capsule", "capsules", "sirop", "syrup", "piure", "pireu",
                }
                distinctive_sa = {w for w in pn_words if w not in _skip_sa}
                if distinctive_sa:
                    # Strict: require ALL distinctive words to be present
                    strict_filtered = [e for e in entries
                                       if all(dw in f"{e['title']} {e['url']}".lower() for dw in distinctive_sa)]
                    if strict_filtered:
                        entries = strict_filtered
                        print(f"[SEARCHANISE] Strict filter: {len(entries)} entries have ALL of {distinctive_sa}")
                    else:
                        # Fallback: require at least most distinctive words
                        best_match_count = 0
                        for e in entries:
                            text = f"{e['title']} {e['url']}".lower()
                            count = sum(1 for dw in distinctive_sa if dw in text)
                            best_match_count = max(best_match_count, count)
                        if best_match_count > 0:
                            threshold = max(1, best_match_count)
                            relaxed_filtered = [e for e in entries
                                                if sum(1 for dw in distinctive_sa
                                                       if dw in f"{e['title']} {e['url']}".lower()) >= threshold]
                            if relaxed_filtered:
                                entries = relaxed_filtered
                                print(f"[SEARCHANISE] Relaxed filter: {len(entries)} entries have {threshold}/{len(distinctive_sa)} of {distinctive_sa}")
                            else:
                                print(f"[SEARCHANISE] No entries match any of {distinctive_sa}")
                                return []

            # Use matcher to pick best entry
            best = None
            if self.ai_matcher and product_name:
                ai_matches = self.ai_matcher.match(product_name, entries, max_results=1)
                if ai_matches:
                    best = ai_matches[0]
            if not best:
                best = entries[0]  # First result from Searchanise is usually best

            # Conflict check — on BOTH url and title
            if product_name:
                if url_has_conflicting_product(best["url"], product_name):
                    print(f"[SEARCHANISE] SKIP conflicting URL: {best['url'][:80]}")
                    return []
                if best.get("title") and url_has_conflicting_product(best["title"], product_name):
                    print(f"[SEARCHANISE] SKIP conflicting TITLE: '{best['title'][:80]}'")
                    return []

            print(f"[SEARCHANISE] Best match: '{best['title']}' → {best['url']}")

            # Extract product images from the matched page
            effective_base = f"https://{site}"
            images = self._extract_product_images(best["url"], effective_base)
            for img_url in images[:max_results]:
                results.append({
                    "image": img_url,
                    "title": best["title"],
                    "source": f"direct:{site}",
                    "product_url": best["url"],
                })

        except Exception as e:
            logger.debug(f"Searchanise API error for {site}: {e}")

        return results

    def search(self, site: str, query: str, max_results: int = 5,
               product_name: str = "") -> list[dict]:
        """
        Search a site directly and return product image URLs.
        If product_name is provided (the original full product name),
        uses AI-powered matching to pick the best product from search results.
        Returns list of {"image": url, "title": ..., "source": "direct:site"}
        """
        results = []
        base_url = f"https://{site}" if not site.startswith("http") else site
        domain = urlparse(base_url).netloc

        # === PHASE 0: Try Searchanise JSON API (most reliable for sites that use it) ===
        sa_results = self._try_searchanise_api(site, query, product_name, max_results)
        if sa_results:
            return sa_results

        # Track if site had results but all were wrong type (for caching)
        _site_had_results_but_all_wrong = False

        # Fetch all search patterns in PARALLEL, then process in order
        from concurrent.futures import ThreadPoolExecutor, as_completed

        def _fetch_pattern(pattern):
            """Fetch one search pattern URL. Returns (pattern, resp) or (pattern, None)."""
            search_url = base_url.rstrip("/") + pattern.format(
                query=requests.utils.quote(query)
            )
            try:
                resp = self.session.get(search_url, timeout=4)
                if resp.status_code == 200:
                    return (pattern, resp)
            except Exception:
                pass
            return (pattern, None)

        pattern_responses = {}
        with ThreadPoolExecutor(max_workers=min(len(self.SEARCH_PATTERNS), 8)) as pool:
            futures = {pool.submit(_fetch_pattern, p): p for p in self.SEARCH_PATTERNS}
            for future in as_completed(futures):
                pattern, resp = future.result()
                pattern_responses[pattern] = resp

        # Process responses in original pattern priority order
        for pattern in self.SEARCH_PATTERNS:
            if len(results) >= max_results:
                break
            resp = pattern_responses.get(pattern)
            if resp is None:
                continue
            try:
                # Update base_url from final URL after any redirects (www vs non-www)
                final_parsed = urlparse(resp.url)
                effective_base = f"{final_parsed.scheme}://{final_parsed.netloc}"

                # Extract product links (filtered by URL patterns + keyword relevance)
                product_links = self._extract_product_links(resp.text, effective_base, query)
                print(f"[SEARCH] Pattern {pattern} → {len(product_links)} links for query='{query}'")
                if product_links:
                    for pl in product_links[:5]:
                        print(f"  link: {pl[:100]}")
                if not product_links:
                    continue  # This search pattern didn't find matching products

                # Build entries with metadata for AI matching
                raw_entries = self._extract_product_entries(resp.text, effective_base)
                url_to_entry = {e["url"]: e for e in raw_entries}
                filtered_entries = []
                for link in product_links:
                    entry = url_to_entry.get(link, {"url": link})
                    filtered_entries.append(entry)

                # === PRE-FILTER: require distinctive product words ===
                # If product name has distinctive words (e.g. "paulista", "india"),
                # require ALL of them. This prevents matching wrong product variants
                # (e.g. "Tierra Selection" instead of "Tierra India").
                if product_name and filtered_entries:
                    pn_words = set(re.findall(r'[a-z]{4,}', product_name.lower()))
                    _skip_pf = set(_RO_EN_MAP.keys()) | set(_RO_EN_MAP.values()) | {
                        "cafea", "coffee", "ceai", "tea", "boabe", "beans",
                        "capsule", "capsules", "sirop", "syrup", "piure", "pireu",
                    }
                    distinctive_pf = {w for w in pn_words if w not in _skip_pf}
                    if distinctive_pf:
                        # Strict: require ALL distinctive words
                        strict_filtered = []
                        for entry in filtered_entries:
                            entry_text = f"{entry.get('title','')} {entry.get('alt','')} {entry.get('all_text','')} {entry['url']}".lower()
                            if all(dw in entry_text for dw in distinctive_pf):
                                strict_filtered.append(entry)
                        if strict_filtered:
                            print(f"[SEARCH] PRE-FILTER: {len(strict_filtered)}/{len(filtered_entries)} entries have ALL distinctive words {distinctive_pf}")
                            filtered_entries = strict_filtered
                        else:
                            # Fallback: keep entries with maximum matching distinctive words
                            best_count = 0
                            for entry in filtered_entries:
                                entry_text = f"{entry.get('title','')} {entry.get('alt','')} {entry.get('all_text','')} {entry['url']}".lower()
                                count = sum(1 for dw in distinctive_pf if dw in entry_text)
                                best_count = max(best_count, count)
                            if best_count > 0:
                                relaxed_filtered = []
                                for entry in filtered_entries:
                                    entry_text = f"{entry.get('title','')} {entry.get('alt','')} {entry.get('all_text','')} {entry['url']}".lower()
                                    count = sum(1 for dw in distinctive_pf if dw in entry_text)
                                    if count >= best_count:
                                        relaxed_filtered.append(entry)
                                filtered_entries = relaxed_filtered
                                print(f"[SEARCH] PRE-FILTER: {len(filtered_entries)} entries have {best_count}/{len(distinctive_pf)} of {distinctive_pf}")
                            else:
                                print(f"[SEARCH] PRE-FILTER: no entries have any of {distinctive_pf}, keeping all {len(filtered_entries)}")

                # === AI-POWERED MATCHING ===
                # Use AI to validate and rank entries against the product we're looking for
                best_entries = []

                if self.ai_matcher and product_name:
                    ai_candidates = [{"title": e.get("title", ""), "url": e["url"],
                                       "alt": e.get("alt", ""),
                                       "all_text": e.get("all_text", "")} for e in filtered_entries[:30]]
                    print(f"[SEARCH] Calling matcher with {len(ai_candidates)} candidates for '{product_name}'")
                    ai_matches = self.ai_matcher.match(product_name, ai_candidates, max_results=3)
                    if ai_matches:
                        best_entries = ai_matches
                        print(f"[SEARCH] Matcher returned: {ai_matches[0].get('title', ai_matches[0].get('url', ''))}")
                        logger.info(f"AI matched '{product_name}' → {ai_matches[0].get('title', ai_matches[0].get('url', ''))}")
                    else:
                        # AI says none of these are the right product → skip this pattern
                        print(f"[SEARCH] Matcher returned EMPTY for '{product_name}' — all {len(ai_candidates)} candidates rejected")
                        logger.debug(f"AI rejected all {len(filtered_entries)} entries for '{product_name}' from pattern {pattern}")
                        continue

                # Fall back to original keyword ranking order (no AI available)
                if not best_entries:
                    # Apply type filter even in fallback path
                    name_type_fb = self.ai_matcher._detect_product_type(product_name) if self.ai_matcher and product_name else None
                    fallback_entries = [{"url": link} for link in product_links[:10]]
                    if name_type_fb:
                        filtered_fb = []
                        for e in fallback_entries:
                            url_text = urlparse(e["url"]).path.replace("-", " ").replace("_", " ")
                            entry_in_map = url_to_entry.get(e["url"], {})
                            cand_text = f"{entry_in_map.get('title','')} {url_text}"
                            cand_type = self.ai_matcher._detect_product_type(cand_text)
                            if cand_type and cand_type != name_type_fb:
                                print(f"[FALLBACK] SKIP type mismatch: {e['url'][:80]} (wanted={name_type_fb}, got={cand_type})")
                                continue
                            filtered_fb.append(e)
                        best_entries = filtered_fb[:3]
                    else:
                        best_entries = fallback_entries[:3]

                if not best_entries:
                    if filtered_entries:
                        _site_had_results_but_all_wrong = True
                    continue

                # CONFLICTING PRODUCT CHECK: skip entries whose URL has wrong flavor/variant
                # Exception: if the entry TITLE closely matches product_name, trust title over URL
                # (some sites have mismatched URLs, e.g. cloned products with old slug)
                if product_name and best_entries:
                    checked_entries = []
                    for entry in best_entries:
                        entry_title = entry.get("title", "")
                        # If title matches product name well, skip URL conflict check
                        # (title is more reliable than URL for product identity)
                        title_trusted = False
                        if entry_title:
                            pn_words = set(re.findall(r'[a-z]{3,}', product_name.lower()))
                            title_words = set(re.findall(r'[a-z]{3,}', entry_title.lower()))
                            _ignore_words = {"ceai", "tea", "cafea", "coffee", "new", "nou", "bio", "eco"}
                            pn_distinctive = pn_words - _ignore_words
                            title_distinctive = title_words - _ignore_words
                            if pn_distinctive and title_distinctive:
                                overlap = pn_distinctive & title_distinctive
                                if len(overlap) >= len(pn_distinctive) * 0.6:
                                    title_trusted = True
                                    print(f"[SEARCH] Title trusted over URL: '{entry_title}' (overlap={overlap})")

                        if not title_trusted:
                            if url_has_conflicting_product(entry["url"], product_name):
                                print(f"[SEARCH] SKIP conflicting product URL: {entry['url'][:100]} for '{product_name}'")
                                continue
                            if entry_title and url_has_conflicting_product(entry_title, product_name):
                                print(f"[SEARCH] SKIP conflicting product TITLE: '{entry_title}' for '{product_name}'")
                                continue
                        checked_entries.append(entry)
                    if checked_entries:
                        best_entries = checked_entries
                    else:
                        # ALL top entries had conflicts — search ALL candidates for ones
                        # containing distinctive product words (e.g. "paulista")
                        name_words_set = set(re.findall(r'[a-z]{4,}', product_name.lower()))
                        _skip = set(_RO_EN_MAP.keys()) | set(_RO_EN_MAP.values()) | {
                            "cafea", "coffee", "ceai", "tea", "boabe", "beans",
                            "capsule", "capsules", "sirop", "syrup", "piure", "pireu",
                        }
                        distinctive = {w for w in name_words_set if w not in _skip}
                        if distinctive:
                            # Look through ALL original filtered_entries (not just matcher's top picks)
                            rescue_entries = []
                            for entry in filtered_entries:
                                entry_text = f"{entry.get('title','')} {entry['url']}".lower()
                                # Entry must contain at least one distinctive word
                                if any(dw in entry_text for dw in distinctive):
                                    entry_url = entry.get("url", "")
                                    entry_title = entry.get("title", "")
                                    # Check if title matches product well (trust title over URL)
                                    rescue_title_trusted = False
                                    if entry_title:
                                        rt_words = set(re.findall(r'[a-z]{3,}', entry_title.lower()))
                                        rt_overlap = (rt_words - {"ceai","tea","cafea","coffee","new","nou"}) & distinctive
                                        if len(rt_overlap) >= 2:
                                            rescue_title_trusted = True
                                    if not rescue_title_trusted:
                                        if entry_url and url_has_conflicting_product(entry_url, product_name):
                                            print(f"[SEARCH] RESCUE SKIP conflict URL: {entry_url}")
                                            continue
                                        if entry_title and url_has_conflicting_product(entry_title, product_name):
                                            print(f"[SEARCH] RESCUE SKIP conflict title: '{entry_title}'")
                                            continue
                                    rescue_entries.append(entry)
                            if rescue_entries:
                                print(f"[SEARCH] RESCUE: found {len(rescue_entries)} entries with distinctive words {distinctive}")
                                best_entries = rescue_entries[:3]
                            else:
                                print(f"[SEARCH] ALL entries had conflicts, no rescue candidates with {distinctive}")
                                best_entries = []  # Clear — don't visit conflicting pages
                                _site_had_results_but_all_wrong = True
                        else:
                            best_entries = []  # No distinctive words to rescue with
                            _site_had_results_but_all_wrong = True

                # Visit the best-matching product page(s)
                for entry in best_entries[:2]:
                    if len(results) >= max_results:
                        break
                    link = entry["url"]
                    print(f"[SEARCH] Visiting product page: {link[:120]} (title='{entry.get('title','')[:60]}')")
                    images = self._extract_product_images(link, effective_base)
                    print(f"[SEARCH] Extracted {len(images)} images from {link[:80]}")
                    for img_url in images[:2]:
                        print(f"  image: {img_url[:120]}")
                    for img_url in images:
                        if len(results) >= max_results:
                            break
                        results.append({
                            "image": img_url,
                            "title": entry.get("title", query),
                            "source": f"direct:{site}",
                            "product_url": link,
                        })

                if results:
                    break  # Found results with this search pattern

            except Exception as e:
                logger.debug(f"Direct scrape error for {site}: {e}")
                continue

        # === EARLY STOP: if search patterns found products but ALL were wrong type/conflicting,
        # skip slug + brand page + Google site:search (they're unlikely to find better results) ===
        if not results and _site_had_results_but_all_wrong:
            print(f"[SEARCH] EARLY STOP: {domain} had results but all conflicted, skipping slug/brand/google phases")
            return results

        # If no results, try constructing a product URL from the full product name (most precise)
        if not results:
            # Use product_name (full name with type words like Pireu/Sirop) for better slug matching
            slug_query = product_name or query
            slug_results = self._try_slug_url(base_url, slug_query, max_results)
            results.extend(slug_results)

        # If still no results, try brand page crawling (broader, less precise)
        if not results:
            brand_results = self._try_brand_page(base_url, query, max_results,
                                                  product_name=product_name)
            results.extend(brand_results)

        # === PHASE 5: Google site: search fallback (universal — works on ANY indexed site) ===
        if not results:
            google_results = self._try_google_site_search(site, query, product_name, max_results)
            results.extend(google_results)

        return results

    def _try_google_site_search(self, site: str, query: str,
                                product_name: str = "", max_results: int = 5) -> list[dict]:
        """
        Universal fallback: search Google with 'site:domain.com product name'.
        Works on ANY indexed site without needing to know its search mechanism.
        This is the last-resort strategy when all direct search methods fail.
        """
        results = []
        base_url = f"https://{site}" if not site.startswith("http") else site
        search_name = product_name or query

        try:
            from googlesearch import search as google_search

            google_query = f"site:{site} {search_name}"
            print(f"[GOOGLE] Searching: {google_query}")

            product_urls = []
            for url in google_search(google_query, num_results=10, lang="ro"):
                # Filter: must be from the target site and look like a product page
                parsed = urlparse(url)
                if not self._domains_match(parsed.netloc, urlparse(base_url).netloc):
                    continue
                if self._is_product_url(url, base_url):
                    product_urls.append(url)
                    print(f"[GOOGLE] Found product URL: {url[:120]}")

            if not product_urls:
                print(f"[GOOGLE] No product URLs found for '{search_name}' on {site}")
                return []

            # Apply AI matching if we have product_name
            best_urls = product_urls[:3]

            if self.ai_matcher and product_name and len(product_urls) > 1:
                candidates = []
                for purl in product_urls[:10]:
                    # Extract title from URL slug
                    slug = urlparse(purl).path.strip('/').split('/')[-1]
                    slug_title = slug.replace('-', ' ').replace('_', ' ').replace('.html', '').replace('.htm', '')
                    candidates.append({
                        "title": slug_title,
                        "url": purl,
                        "alt": slug_title,
                        "all_text": slug_title,
                    })
                ai_matches = self.ai_matcher.match(product_name, candidates, max_results=3)
                if ai_matches:
                    best_urls = [m["url"] for m in ai_matches]
                    print(f"[GOOGLE] AI matched: {best_urls[0][:100]}")

            # Apply conflict check
            checked_urls = []
            for url in best_urls:
                if product_name and url_has_conflicting_product(url, product_name):
                    print(f"[GOOGLE] SKIP conflicting: {url[:100]}")
                    continue
                checked_urls.append(url)
            if not checked_urls:
                checked_urls = best_urls[:1]  # Keep at least one

            # Visit product pages and extract images
            for purl in checked_urls[:2]:
                if len(results) >= max_results:
                    break
                effective_base = f"{urlparse(purl).scheme}://{urlparse(purl).netloc}"
                print(f"[GOOGLE] Visiting: {purl[:120]}")

                # Check page title for conflict
                try:
                    resp = self.session.get(purl, timeout=6)
                    if resp.status_code != 200:
                        continue
                    title_m = re.search(r'<title[^>]*>([^<]+)', resp.text, re.IGNORECASE)
                    page_title = title_m.group(1).strip() if title_m else ""
                    if page_title and product_name and url_has_conflicting_product(page_title, product_name):
                        print(f"[GOOGLE] SKIP conflicting page title: '{page_title[:80]}'")
                        continue
                except Exception:
                    continue

                images = self._extract_product_images(purl, effective_base)
                for img_url in images:
                    if len(results) >= max_results:
                        break
                    results.append({
                        "image": img_url,
                        "title": product_name or query,
                        "source": f"direct:{site}",
                        "product_url": purl,
                    })

            if results:
                print(f"[GOOGLE] Found {len(results)} results via Google site: search")

        except ImportError:
            print("[GOOGLE] googlesearch-python not installed, skipping Google fallback")
        except Exception as e:
            print(f"[GOOGLE] Google site: search failed: {e}")

        return results

    def _try_brand_page(self, base_url: str, query: str, max_results: int = 5,
                        product_name: str = "") -> list[dict]:
        """
        Try to find products by crawling the brand/category page.
        Many sites have /brands/monin or /monin pages listing all brand products.
        We crawl those and find the matching product.
        """
        results = []
        site = urlparse(base_url).netloc

        # Extract brand name (usually the first distinctive keyword)
        cleaned = clean_product_query(query)
        skip_words = set(_RO_EN_MAP.keys()) | {"de", "cu", "si", "din", "la", "pt", "pentru"}
        key_words = [w for w in cleaned.split() if w.lower() not in skip_words and len(w) > 1]

        if not key_words:
            return []

        # The brand is typically the first key word or first two words
        brand_candidates = [key_words[0].lower()]
        if len(key_words) >= 2:
            brand_candidates.append(f"{key_words[0]}-{key_words[1]}".lower())

        # Brand page URL patterns to try
        brand_patterns = [
            "/brands/{brand}",
            "/{brand}",
            "/brand/{brand}",
            "/producator/{brand}",
            "/manufacturer/{brand}",
        ]

        for brand in brand_candidates:
            if results:
                break
            for pattern in brand_patterns:
                if results:
                    break
                url = base_url.rstrip("/") + pattern.format(brand=brand)
                try:
                    # Check brand page cache first
                    cache_key = url
                    cached = self._brand_page_cache.get(cache_key)
                    if cached is not None:
                        entries, effective_base = cached
                        print(f"[BRAND] Brand page {url} → {len(entries)} entries (CACHED)")
                    else:
                        resp = self.session.get(url, timeout=6)
                        if resp.status_code != 200:
                            self._brand_page_cache[cache_key] = ([], "")
                            continue

                        # Update base from final redirect URL
                        final_parsed = urlparse(resp.url)
                        effective_base = f"{final_parsed.scheme}://{final_parsed.netloc}"

                        # Check this page has product links (not just a generic 404)
                        entries = self._extract_product_entries(resp.text, effective_base)
                        self._brand_page_cache[cache_key] = (entries, effective_base)
                        print(f"[BRAND] Brand page {url} → {len(entries)} entries")
                    for e in entries[:5]:
                        print(f"  entry: {e.get('title','(no title)')[:60]} | {e['url'][:80]}")
                    if len(entries) < 2:
                        continue  # Too few = probably not a brand listing page

                    # === AI-POWERED MATCHING on brand page ===
                    match_name = product_name or query
                    best = None

                    if self.ai_matcher and product_name:
                        ai_candidates = [{"title": e.get("title", ""), "url": e["url"],
                                           "alt": e.get("alt", "")} for e in entries[:40]]
                        ai_matches = self.ai_matcher.match(product_name, ai_candidates, max_results=1)
                        if ai_matches:
                            best = ai_matches[0]
                            logger.info(f"AI brand match: '{product_name}' → {best.get('title', best['url'])}")

                    # Fall back to keyword+synonym ranking
                    if not best:
                        entries = self._rank_entries_by_relevance(entries, query)
                        # Also try translated query for ranking
                        translated_words = []
                        for w in query.split():
                            en = _RO_EN_MAP.get(w.lower())
                            translated_words.append(en if en else w)
                        translated_q = " ".join(translated_words)
                        if translated_q != query:
                            entries_t = self._rank_entries_by_relevance(
                                self._extract_product_entries(resp.text, effective_base),
                                translated_q
                            )
                            url_to_entry = {e["url"]: e for e in entries}
                            for e in entries_t:
                                if e["url"] not in url_to_entry:
                                    entries.append(e)

                        # Filter to entries with relevance > 0 using synonym-aware scoring
                        qkw = [w.lower() for w in query.split() if len(w) >= 2]
                        qkw_t = [w.lower() for w in translated_q.split() if len(w) >= 2]
                        all_kw = list(set(qkw + qkw_t))
                        kw_variants = [get_word_variants(k) for k in all_kw]

                        def _entry_score(e):
                            tl = e.get("title", "").lower()
                            al = e.get("alt", "").lower()
                            at = e.get("all_text", "").lower()
                            un = urlparse(e["url"]).path.lower().replace("-", " ")
                            combined = f"{tl} {al} {un} {at}"
                            entry_words = set(re.findall(r'[a-zA-ZăâîșțéèêëüöïôàáùúñçÀ-ÿ]+', combined))
                            s = 0
                            for k, variants in zip(all_kw, kw_variants):
                                if k in tl: s += 3
                                elif k in al: s += 2
                                elif k in un: s += 1.5
                                elif k in at: s += 0.5
                                else:
                                    matched = variants & entry_words
                                    if matched:
                                        if any(v in tl for v in matched): s += 2.5
                                        elif any(v in un for v in matched): s += 1.2
                                        else: s += 0.4
                            return s / len(all_kw) if all_kw else 0

                        relevant = [e for e in entries if _entry_score(e) > 0]
                        if not relevant:
                            continue

                        # TYPE FILTER: if product_name specifies a type (pireu/sirop/ceai),
                        # reject candidates with a DIFFERENT type (same logic as _local_match)
                        name_type = self.ai_matcher._detect_product_type(match_name) if self.ai_matcher else None
                        if name_type:
                            type_filtered = []
                            for e in relevant:
                                cand_text = f"{e.get('title','')} {urlparse(e['url']).path.replace('-',' ')}"
                                cand_type = self.ai_matcher._detect_product_type(cand_text)
                                if cand_type and cand_type != name_type:
                                    print(f"[BRAND] SKIP type mismatch: '{e.get('title','')}' (wanted={name_type}, got={cand_type})")
                                    continue
                                type_filtered.append(e)
                            if type_filtered:
                                relevant = type_filtered
                            # If ALL were rejected, keep original list (better something than nothing)

                        best = max(relevant, key=_entry_score)

                    logger.debug(f"Brand page match: {best.get('title', '')} -> {best['url']}")

                    # CONFLICTING PRODUCT CHECK on brand page match
                    if match_name and url_has_conflicting_product(best["url"], match_name):
                        print(f"[BRAND] SKIP conflicting product in URL: {best['url'][:100]} for '{match_name}'")
                        continue
                    if match_name and best.get("title") and url_has_conflicting_product(best["title"], match_name):
                        print(f"[BRAND] SKIP conflicting product in TITLE: '{best['title'][:80]}' for '{match_name}'")
                        continue

                    print(f"[BRAND] Visiting product page: {best['url'][:100]} (title='{best.get('title','')[:60]}')")
                    images = self._extract_product_images(best["url"], effective_base)
                    for img_url in images:
                        if len(results) >= max_results:
                            break
                        results.append({
                            "image": img_url,
                            "title": query,
                            "source": f"direct:{site}",
                        })
                    if results:
                        break
                except Exception as e:
                    logger.debug(f"Brand page crawl error: {e}")
                    continue

        return results

    def _try_slug_url(self, base_url: str, query: str, max_results: int = 5) -> list[dict]:
        """
        Try to construct a product URL by slugifying the query.
        Many e-commerce sites use URL slugs like /product-name-here or
        /product-name-here.html (CS-Cart, Magento).
        Generates multiple slug variations and tests them.
        """
        results = []
        site = urlparse(base_url).netloc

        # Extract volume/weight BEFORE cleaning (for slug suffix)
        vol_match = re.search(r'\b(\d+(?:[.,]\d+)?)\s*(l|ml|gr|g|kg|cl)\b', query, re.IGNORECASE)
        volume_suffix = ""
        if vol_match:
            vol_num = vol_match.group(1).replace(",", ".")
            vol_unit = vol_match.group(2).lower()
            # Normalize: "gr" → "g", "1.0l" → "1l"
            if vol_unit == "gr":
                vol_unit = "g"
            volume_suffix = f"{vol_num}{vol_unit}".replace(".0", "")

        # Clean the query: remove weights, packaging noise, parenthetical hints
        cleaned = clean_product_query(query)
        # Remove parenthetical content like "(mere verzi)" — it's a hint, not part of the slug
        cleaned = re.sub(r'\([^)]*\)', '', cleaned).strip()
        cleaned = re.sub(r'\s+', ' ', cleaned)

        def _ro_word_forms(word: str) -> set[str]:
            """Generate common Romanian morphological variants of a word.
            Handles plural/singular, articulated/unarticulated forms common in URLs.
            e.g. mandarine → {mandarine, mandarina, mandarin}
                 capsuni → {capsuni, capsuna, capsune}
                 banane → {banane, banana, banan}
            """
            w = word.lower()
            forms = {w}
            # -ine → -ina, -in (mandarine → mandarina, mandarin)
            if w.endswith("ine"):
                forms.add(w[:-1] + "a")   # mandarine → mandarina
                forms.add(w[:-2])          # mandarine → mandarin
            # -une → -una, -un (capsune → capsuna)
            if w.endswith("une"):
                forms.add(w[:-1] + "a")
                forms.add(w[:-2])
            # -ane → -ana, -an (banane → banana)
            if w.endswith("ane"):
                forms.add(w[:-1] + "a")
                forms.add(w[:-2])
            # -ale → -ala (portocale → portocala)
            if w.endswith("ale"):
                forms.add(w[:-1] + "a")
            # -uri → -ura (piure stays, but fructuri → fructura)
            if w.endswith("uri"):
                forms.add(w[:-1] + "a")
            # -e → -a (generic feminine: lamaie→lamaia is different, but piure→piura)
            if w.endswith("e") and len(w) > 3 and not w.endswith(("ine", "une", "ane", "ale")):
                forms.add(w[:-1] + "a")
            # -a → -e (reverse: banana → banane, mandarina → mandarine)
            if w.endswith("a") and len(w) > 3:
                forms.add(w[:-1] + "e")
            # -i → -a, -e (capsuni → capsuna, capsune)
            if w.endswith("i") and len(w) > 3:
                forms.add(w[:-1] + "a")
                forms.add(w[:-1] + "e")
            return forms

        def make_slug(text: str) -> str:
            s = re.sub(r'[^a-z0-9\s-]', '', text.lower())
            s = re.sub(r'\s+', '-', s.strip())
            return re.sub(r'-+', '-', s).strip('-')

        words = cleaned.split()

        # Classify words: brand, product_type, flavor
        product_type_en = {"puree", "syrup", "tea", "coffee", "infusion", "juice",
                           "sauce", "cream", "jam", "compote", "butter", "honey",
                           "chocolate", "liqueur", "concentrate", "powder"}
        filler = {"de", "cu", "si", "din", "la", "pt", "pentru", "sau", "mic", "mare", "mediu", "mini"}
        # Descriptor words: color/attribute modifiers that aren't product identifiers
        _slug_noise = {"negru", "alb", "rosu", "verde", "new", "nou", "vechi", "special",
                       "premium", "clasic", "original", "traditional", "bio", "eco"}

        brand_words = []
        type_words = []   # (original, EN)
        flavor_words = []  # (original, EN)
        descriptor_words = []  # words like "negru" that may or may not be useful
        for w in words:
            lower = w.lower()
            if lower in filler or len(w) <= 1 or re.match(r'^\d', w):
                continue
            if lower in _slug_noise:
                descriptor_words.append(w)
                continue
            en = normalize_to_english(lower)
            if en and en.lower() in product_type_en:
                type_words.append((w, en))
            elif en and en.lower() != lower:
                flavor_words.append((w, en))
            else:
                brand_words.append(w)

        # Build slug variants — keep it small (max ~20 slugs × 2 URL patterns = ~40 requests)
        slugs = []  # ordered by likelihood

        def _top_forms(word_orig: str, word_en: str, limit: int = 4) -> list[str]:
            """Get the most likely slug forms: original, EN, and RO morphological variants."""
            forms = []
            seen = set()
            for candidate in [word_orig.lower(), word_en.lower()]:
                if candidate not in seen:
                    seen.add(candidate)
                    forms.append(candidate)
            # Add RO morphological variants of the original word
            for f in _ro_word_forms(word_orig):
                if f not in seen and re.match(r'^[a-z]+$', f):
                    seen.add(f)
                    forms.append(f)
            # Add key synonym variants (e.g. piure for puree)
            for v in get_word_variants(word_en):
                if v not in seen and re.match(r'^[a-z]+$', v) and len(v) >= 3:
                    seen.add(v)
                    forms.append(v)
                    # Also add RO forms of this variant
                    for rf in _ro_word_forms(v):
                        if rf not in seen and re.match(r'^[a-z]+$', rf):
                            seen.add(rf)
                            forms.append(rf)
            return forms[:limit]

        brand_str = " ".join(brand_words)
        t_forms = []
        for tw in type_words:
            t_forms = _top_forms(tw[0], tw[1], limit=4)
        f_forms = []
        for fw in flavor_words:
            f_forms = _top_forms(fw[0], fw[1], limit=4)

        # Variant 1: brand-flavor-type (monin-mandarina-piure)
        if brand_words and f_forms and t_forms:
            for ff in f_forms:
                for tf in t_forms:
                    s = make_slug(f"{brand_str} {ff} {tf}")
                    if s and s not in slugs:
                        slugs.append(s)

        # Variant 2: brand-flavor (monin-mandarina)
        if brand_words and f_forms:
            for ff in f_forms:
                s = make_slug(f"{brand_str} {ff}")
                if s and s not in slugs:
                    slugs.append(s)

        # Variant 2b: brand-type (when flavor is in brand_words, e.g. MONIN Green Apple + piure)
        if brand_words and t_forms and not f_forms:
            for tf in t_forms:
                s = make_slug(f"{brand_str} {tf}")
                if s and s not in slugs:
                    slugs.append(s)

        # Variant 3: type-brand-flavor (piure-monin-mandarina)
        if brand_words and t_forms:
            flavor_part = f_forms[0] if f_forms else ""
            for tf in t_forms[:2]:
                full = f"{tf} {brand_str} {flavor_part}".strip()
                s = make_slug(full)
                if s and s not in slugs:
                    slugs.append(s)

        # Variant 4: Full cleaned name as-is
        s = make_slug(cleaned)
        if s and s not in slugs:
            slugs.append(s)

        # Variant 5: Translated full name
        translated_words = []
        for w in words:
            en = _RO_EN_MAP.get(w.lower())
            translated_words.append(en if en else w)
        s = make_slug(" ".join(translated_words))
        if s and s not in slugs:
            slugs.append(s)

        # Variant 6: Just key words (brand + flavor, skip filler + type)
        # Skip if it would be brand-only (too generic, matches category pages)
        key_words = [w for w in words if w.lower() not in filler
                     and w.lower() not in set(_RO_EN_MAP.keys()) and len(w) > 1]
        if key_words and len(key_words) >= 2:
            s = make_slug(" ".join(key_words))
            if s and s not in slugs:
                slugs.append(s)

        # Variant 7: Multi-word type expansions with volume — PRIORITIZED
        # Many sites use full descriptive names like "monin-cherry-piure-de-fructe-1l"
        # These are inserted at position 0 (highest priority) since they're the most specific
        _type_expansions = {
            "puree": ["piure de fructe", "fruit puree", "piure fructe"],
            "syrup": ["sirop de", "sirop"],
            "sauce": ["sos de", "sos"],
            "juice": ["suc de", "suc natural"],
            "jam": ["gem de", "dulceata de"],
            "concentrate": ["concentrat de"],
        }
        expansion_slugs = []
        if brand_words and f_forms and type_words:
            en_type = type_words[0][1].lower()  # e.g. "puree"
            expansions = _type_expansions.get(en_type, [])
            for exp in expansions:
                for ff in f_forms[:2]:  # Limit to top 2 flavor forms
                    # With volume first (most specific): monin-cherry-piure-de-fructe-1l
                    if volume_suffix:
                        sv = make_slug(f"{brand_str} {ff} {exp} {volume_suffix}")
                        if sv and sv not in expansion_slugs:
                            expansion_slugs.append(sv)
                    # Without volume: monin-cherry-piure-de-fructe
                    s = make_slug(f"{brand_str} {ff} {exp}")
                    if s and s not in expansion_slugs:
                        expansion_slugs.append(s)
        # Insert expansion slugs at the front (most likely to match)
        slugs = expansion_slugs + [s for s in slugs if s not in expansion_slugs]

        # Variant 8: Add volume suffix to top basic slugs
        if volume_suffix:
            vol_slugs = []
            for s in slugs[:8]:
                sv = f"{s}-{make_slug(volume_suffix)}"
                if sv not in slugs and sv not in vol_slugs:
                    vol_slugs.append(sv)
            slugs.extend(vol_slugs)

        # Variant 9: Condensed brand slugs
        # Detect multi-word brand names where adjacent words should be joined
        # e.g. "Tea Tales" → "teatales", "Dr Oetker" → "droetker"
        # Strategy: try joining ANY consecutive 2-3 word sequence as brand prefix,
        # then append remaining words as hyphenated product description.

        # Reconstruct the word order from the original cleaned query
        _ordered_words = [w for w in words if w.lower() not in filler
                          and w.lower() not in _slug_noise
                          and not re.match(r'^\d', w.lower())]

        # Try consecutive 2-word spans that include at least one type/brand word
        # (skip 3-word spans to reduce URL count — 2-word condensed is most common)
        _condensed_combos = []  # (condensed_str, start_idx, span_len)
        type_word_lower = set(tw[0].lower() for tw in type_words)
        brand_word_lower = set(w.lower() for w in brand_words)
        for start in range(len(_ordered_words) - 1):
            w1 = _ordered_words[start].lower()
            w2 = _ordered_words[start + 1].lower()
            # Only combine if one is type/brand and the other is brand
            # (avoids combining random adjacent words)
            is_relevant = ((w1 in type_word_lower and w2 in brand_word_lower) or
                           (w1 in brand_word_lower and w2 in brand_word_lower) or
                           (w1 in type_word_lower and w2 not in type_word_lower))
            if is_relevant:
                condensed = w1 + w2
                if len(condensed) >= 4:
                    _condensed_combos.append((condensed, start, 2))

        # Also try type_word + adjacent brand_word
        if type_words and brand_words:
            for tw in type_words:
                cb = tw[0].lower() + brand_words[0].lower()
                if len(cb) >= 4:
                    # Find positions in _ordered_words
                    try:
                        tw_idx = next(i for i, w in enumerate(_ordered_words) if w.lower() == tw[0].lower())
                        bw_idx = next(i for i, w in enumerate(_ordered_words) if w.lower() == brand_words[0].lower())
                        if bw_idx == tw_idx + 1:
                            _condensed_combos.append((cb, tw_idx, 2))
                    except StopIteration:
                        pass

        # Deduplicate combos by condensed string
        seen_condensed = set()
        unique_combos = []
        for combo in _condensed_combos:
            if combo[0] not in seen_condensed:
                seen_condensed.add(combo[0])
                unique_combos.append(combo)

        for condensed, start_idx, span_len in unique_combos:
            # Remaining words = everything NOT in the condensed span
            remaining_parts = []
            for i, w in enumerate(_ordered_words):
                if start_idx <= i < start_idx + span_len:
                    continue  # Skip words that form the condensed brand
                lower = w.lower()
                en = normalize_to_english(lower)
                if en and en.lower() in product_type_en:
                    continue  # Skip type words in remainder (ceai, infuzie)
                if en and en.lower() != lower:
                    remaining_parts.append(en.lower())
                else:
                    remaining_parts.append(lower)
            remaining_str = "-".join(remaining_parts) if remaining_parts else ""

            if remaining_str:
                sv = f"{condensed}-{remaining_str}"
            else:
                sv = condensed
            # With volume
            if volume_suffix:
                sv_vol = f"{sv}-{make_slug(volume_suffix)}"
                if sv_vol not in slugs:
                    slugs.append(sv_vol)
            if sv not in slugs:
                slugs.append(sv)

        # Variant 10: Type-specific slug aliases (HIGH PRIORITY)
        # Some sites use different terminology, e.g. "infuzie" → "looseleaf"
        # These are inserted near the front because they target specific site conventions
        _type_slug_aliases = {
            "infusion": ["looseleaf", "herbal-tea", "herbal"],
        }
        alias_slugs = []
        if type_words:
            # Only use the most plausible 2-word condensed brands
            # Filter to combos that include at least one type word (likely brand = type+name)
            type_word_set = set(tw[0].lower() for tw in type_words)
            best_combos = []
            for cb_str, cb_start, cb_span in unique_combos:
                if cb_span != 2:
                    continue
                # Check if this combo includes a type word followed by a non-type word
                span_words = [_ordered_words[cb_start+i].lower() for i in range(cb_span)]
                has_type = any(w in type_word_set for w in span_words)
                has_non_type = any(w not in type_word_set for w in span_words)
                if has_type and has_non_type:
                    best_combos.append((cb_str, cb_start, cb_span))

            for en_type_key in set(tw[1].lower() for tw in type_words):
                aliases = _type_slug_aliases.get(en_type_key, [])
                if not aliases:
                    continue
                for cb_str, cb_start, cb_span in best_combos[:2]:  # Max 2 brand combos
                    # Get remaining words (skip condensed span AND type words)
                    product_parts = []
                    for i, w in enumerate(_ordered_words):
                        if cb_start <= i < cb_start + cb_span:
                            continue
                        lower = w.lower()
                        en = normalize_to_english(lower)
                        if en and en.lower() in product_type_en:
                            continue
                        if en and en.lower() != lower:
                            product_parts.append(en.lower())
                        else:
                            product_parts.append(lower)

                    for alias in aliases[:2]:  # Max 2 aliases per type
                        parts = [cb_str, alias] + product_parts
                        sv = "-".join(parts)
                        if volume_suffix:
                            sv_vol = f"{sv}-{make_slug(volume_suffix)}"
                            if sv_vol not in alias_slugs:
                                alias_slugs.append(sv_vol)
                        if sv not in alias_slugs:
                            alias_slugs.append(sv)

        # Insert alias slugs after expansion slugs but before other variants
        # They target specific site conventions so should be tried early
        slugs = slugs[:20] + [s for s in alias_slugs if s not in slugs] + slugs[20:]

        # Limit total slugs to avoid too many HTTP requests
        slugs = slugs[:30]

        # URL suffixes to try (no prefix — product slugs are almost always at root)
        suffixes = ["", ".html"]  # CS-Cart, Magento often use .html

        # Build all candidate URLs (slug × suffix)
        candidate_urls = []
        for slug in slugs:
            for suffix in suffixes:
                url = base_url.rstrip("/") + f"/{slug}" + suffix
                candidate_urls.append((url, slug))

        # --- PARALLEL slug probing with ThreadPoolExecutor ---
        from concurrent.futures import ThreadPoolExecutor, as_completed
        import threading

        found_result = threading.Event()  # Signal other threads to stop early
        ai_matcher = self.ai_matcher

        def _probe_slug(url_slug_pair):
            """Probe a single slug URL. Returns (url, slug, resp) or None."""
            url, slug = url_slug_pair
            if found_result.is_set():
                return None
            try:
                resp = self.session.get(url, timeout=8)
                if resp.status_code != 200:
                    return None

                # Quick product page check
                text_lower = resp.text.lower()
                is_product = (
                    '"@type":"Product"' in resp.text or
                    '"@type": "Product"' in resp.text or
                    'og:image' in resp.text or
                    'data-zoom-image' in resp.text or
                    'add-to-cart' in text_lower or
                    'addtocart' in text_lower
                )
                if is_product and '-' not in slug:
                    has_jsonld = ('"@type":"Product"' in resp.text or
                                  '"@type": "Product"' in resp.text)
                    if not has_jsonld:
                        return None
                if not is_product:
                    return None

                return (url, slug, resp)
            except Exception:
                return None

        # Probe slugs in parallel (8 workers — balanced speed vs politeness)
        valid_hits = []
        with ThreadPoolExecutor(max_workers=8) as pool:
            futures = {pool.submit(_probe_slug, pair): pair for pair in candidate_urls}
            for future in as_completed(futures):
                if found_result.is_set():
                    break
                hit = future.result()
                if hit:
                    valid_hits.append(hit)
                    found_result.set()  # Signal other threads to stop

        # Process valid hits (in slug priority order)
        slug_order = {url: i for i, (url, _) in enumerate(candidate_urls)}
        valid_hits.sort(key=lambda h: slug_order.get(h[0], 999))

        for url, slug, resp in valid_hits:
            if len(results) >= max_results:
                break

            # TYPE CHECK: verify the found page matches our product type
            if ai_matcher:
                name_type_slug = ai_matcher._detect_product_type(query)
                if name_type_slug:
                    title_match = re.search(r'<title[^>]*>([^<]+)', resp.text, re.IGNORECASE)
                    h1_match = re.search(r'<h1[^>]*>([^<]+)', resp.text, re.IGNORECASE)
                    page_title = (title_match.group(1) if title_match else "") + " " + (h1_match.group(1) if h1_match else "")
                    page_text = f"{page_title} {slug}"
                    page_type = ai_matcher._detect_product_type(page_text)
                    if page_type and page_type != name_type_slug:
                        print(f"[SLUG] SKIP type mismatch: {url} (wanted={name_type_slug}, got={page_type}, title='{page_title.strip()[:60]}')")
                        continue

            # CONFLICTING PRODUCT CHECK on slug URL and page title
            if url_has_conflicting_product(url, query):
                print(f"[SLUG] SKIP conflicting product in URL: {url[:100]} for '{query}'")
                continue
            title_match_chk = re.search(r'<title[^>]*>([^<]+)', resp.text, re.IGNORECASE)
            h1_match_chk = re.search(r'<h1[^>]*>([^<]+)', resp.text, re.IGNORECASE)
            page_title_chk = (title_match_chk.group(1) if title_match_chk else "") + " " + (h1_match_chk.group(1) if h1_match_chk else "")
            if page_title_chk.strip() and url_has_conflicting_product(page_title_chk, query):
                print(f"[SLUG] SKIP conflicting product in page title: '{page_title_chk.strip()[:80]}' for '{query}'")
                continue

            effective_base = f"{urlparse(resp.url).scheme}://{urlparse(resp.url).netloc}"
            print(f"[SLUG] Visiting product page: {url[:100]}")
            images = self._extract_product_images(url, effective_base)
            for img_url in images:
                if len(results) >= max_results:
                    break
                results.append({
                    "image": img_url,
                    "title": query,
                    "source": f"direct:{site}",
                    "product_url": url,
                })
            if results:
                print(f"[SLUG] Found matching product: {url}")
                return results

        return results

    def _extract_product_entries(self, html: str, base_url: str) -> list[dict]:
        """
        Extract product entries from search results HTML.
        Each entry: {"url": str, "title": str, "alt": str, "all_text": str}

        Uses multiple strategies for maximum compatibility:
        1. Standard <a> tag scanning (Magento, Shopify, WooCommerce)
        2. BeautifulSoup structured parsing (product containers with classes)
        3. CS-Cart product_id link resolution (dispatch=products.view)
        4. JSON-LD structured data (schema.org/Product)
        """
        entries = []
        seen_urls = set()

        # ── Strategy 1: Regex-based <a> tag scanning (fast, works on most sites) ──
        for m in re.finditer(
            r'<a\s[^>]*href=["\']([^"\']+)["\'][^>]*>(.*?)</a>',
            html, re.IGNORECASE | re.DOTALL
        ):
            href = m.group(1)
            inner_html = m.group(2)
            full_tag = m.group(0)

            url = self._make_absolute(href, base_url)
            if not url or url in seen_urls or not self._is_product_url(url, base_url):
                continue

            title = ""
            title_m = re.search(r'title=["\']([^"\']+)["\']', full_tag)
            if title_m:
                title = title_m.group(1).strip()

            alt_texts = re.findall(r'alt=["\']([^"\']+)["\']', inner_html, re.IGNORECASE)

            link_text = re.sub(r'<[^>]+>', ' ', inner_html).strip()
            link_text = re.sub(r'\s+', ' ', link_text).strip()

            if not title:
                for alt in alt_texts:
                    if len(alt) > 3:
                        title = alt
                        break
            if not title and len(link_text) > 3:
                title = link_text

            start_pos = m.start()
            end_pos = m.end()
            context = html[max(0, start_pos - 500):min(len(html), end_pos + 500)]

            context_alts = re.findall(r'<img[^>]+alt=["\']([^"\']+)["\']', context, re.IGNORECASE)

            name_matches = re.findall(
                r'class=["\'][^"\']*product[- ]?name[^"\']*["\'][^>]*>([^<]+)<',
                context, re.IGNORECASE
            )

            all_parts = [title] + alt_texts + context_alts + name_matches
            all_text = " ".join(filter(None, all_parts))
            all_text = " ".join(dict.fromkeys(all_text.split()))

            if not title and not all_text:
                continue

            seen_urls.add(url)
            entries.append({
                "url": url,
                "title": title,
                "alt": " ".join(alt_texts + context_alts),
                "all_text": all_text,
            })

        # ── Strategy 2: BeautifulSoup product container parsing ──
        # Finds product cards/blocks by common CSS class patterns across platforms
        if len(entries) < 3:
            try:
                soup = BeautifulSoup(html, 'html.parser')
                # Common product container class patterns across e-commerce platforms
                product_selectors = [
                    '[class*="product-item"]', '[class*="product_item"]',
                    '[class*="product-card"]', '[class*="product_card"]',
                    '[class*="product-block"]', '[class*="product_block"]',
                    '[class*="product-layout"]', '[class*="product-grid"]',
                    '[class*="productItem"]', '[class*="productCard"]',
                    '.ty-grid-list__item',  # CS-Cart specific
                    '.ty-search-results__item',  # CS-Cart search results
                    'li[class*="product"]',
                    'div[class*="product"][class*="list"]',
                ]
                product_containers = []
                for sel in product_selectors:
                    try:
                        found = soup.select(sel)
                        product_containers.extend(found)
                    except Exception:
                        continue

                for container in product_containers:
                    # Find the main product link
                    link = container.find('a', href=True)
                    if not link:
                        continue
                    href = link.get('href', '')
                    url = self._make_absolute(href, base_url)
                    if not url or url in seen_urls:
                        continue
                    # Accept if it's a product URL OR has product_id param
                    if not self._is_product_url(url, base_url):
                        if 'product_id' not in href and 'dispatch=products' not in href:
                            continue

                    # Extract product info from the container
                    title = link.get('title', '')
                    img = container.find('img')
                    alt_text = img.get('alt', '') if img else ''
                    # Get all text from the container
                    container_text = container.get_text(separator=' ', strip=True)

                    if not title and alt_text and len(alt_text) > 3:
                        title = alt_text
                    if not title:
                        # Look for product name element
                        name_el = container.select_one(
                            '[class*="product-name"], [class*="product_name"], '
                            '[class*="product-title"], [class*="ty-grid-list__item-name"], '
                            'h2, h3, h4'
                        )
                        if name_el:
                            title = name_el.get_text(strip=True)

                    if not title and len(container_text) > 5:
                        # Use first 80 chars of container text as title
                        title = container_text[:80]

                    if title or container_text:
                        seen_urls.add(url)
                        entries.append({
                            "url": url,
                            "title": title,
                            "alt": alt_text,
                            "all_text": f"{title} {alt_text} {container_text[:200]}",
                        })
            except Exception as e:
                print(f"[EXTRACT] BeautifulSoup strategy failed: {e}")

        # ── Strategy 3: CS-Cart product_id resolution ──
        # CS-Cart puts product links as ?dispatch=products.view&product_id=XXX
        # or in onclick/data attributes. Resolve these to real slug URLs.
        if len(entries) < 3:
            product_ids = set()
            # Find product_id references in HTML
            for pid_match in re.finditer(r'product_id[=:][\s"\']*(\d{3,})', html):
                product_ids.add(pid_match.group(1))
            # Also check dispatch=products.view links
            for pid_match in re.finditer(r'dispatch=products\.view[^"\']*product_id=(\d+)', html):
                product_ids.add(pid_match.group(1))
            # Also from quick_view links
            for pid_match in re.finditer(r'dispatch=products\.quick_view[^"\']*product_id=(\d+)', html):
                product_ids.add(pid_match.group(1))

            if product_ids:
                print(f"[EXTRACT] Found {len(product_ids)} CS-Cart product_ids: {list(product_ids)[:10]}")
                # Resolve each product_id to its slug URL via redirect
                from concurrent.futures import ThreadPoolExecutor, as_completed

                def _resolve_product_id(pid):
                    """Resolve CS-Cart product_id to slug URL via 301/302 redirect."""
                    resolve_url = f"{base_url.rstrip('/')}/index.php?dispatch=products.view&product_id={pid}"
                    try:
                        resp = self.session.head(resolve_url, timeout=8, allow_redirects=True)
                        final_url = resp.url
                        if resp.status_code == 200 and final_url != resolve_url:
                            return (pid, final_url)
                        # Try GET if HEAD doesn't redirect
                        resp = self.session.get(resolve_url, timeout=8, allow_redirects=True)
                        if resp.status_code == 200:
                            final_url = resp.url
                            # Extract title from the product page
                            title_m = re.search(r'<title[^>]*>([^<]+)', resp.text, re.IGNORECASE)
                            title = title_m.group(1).strip() if title_m else ""
                            # Clean title: remove site name suffix
                            title = re.split(r'\s*[-|–]\s*', title)[0].strip()
                            return (pid, final_url, title, resp.text)
                    except Exception:
                        pass
                    return (pid, None)

                with ThreadPoolExecutor(max_workers=min(len(product_ids), 4)) as pool:
                    futures = [pool.submit(_resolve_product_id, pid) for pid in list(product_ids)[:15]]
                    for future in as_completed(futures):
                        result = future.result()
                        if result and len(result) >= 2 and result[1]:
                            pid = result[0]
                            resolved_url = result[1]
                            title = result[2] if len(result) > 2 else ""
                            page_html = result[3] if len(result) > 3 else ""
                            if resolved_url not in seen_urls:
                                seen_urls.add(resolved_url)
                                # Try to get product info from the resolved page
                                alt_text = ""
                                if page_html:
                                    og_title = re.search(r'<meta[^>]+property=["\']og:title["\'][^>]+content=["\']([^"\']+)', page_html, re.IGNORECASE)
                                    if og_title:
                                        title = og_title.group(1).strip() or title
                                    img_alt = re.search(r'<img[^>]+alt=["\']([^"\']{5,})["\']', page_html, re.IGNORECASE)
                                    if img_alt:
                                        alt_text = img_alt.group(1).strip()
                                print(f"[EXTRACT] CS-Cart pid={pid} → {resolved_url[:100]} (title='{title[:60]}')")
                                entries.append({
                                    "url": resolved_url,
                                    "title": title,
                                    "alt": alt_text,
                                    "all_text": f"{title} {alt_text}",
                                })

        # ── Strategy 4: JSON-LD structured data ──
        if len(entries) < 3:
            try:
                import extruct
                data = extruct.extract(html, base_url=base_url, syntaxes=['json-ld', 'opengraph'])
                for item in data.get('json-ld', []):
                    if isinstance(item, dict):
                        items_to_check = [item]
                        # Handle @graph arrays
                        if '@graph' in item:
                            items_to_check = item['@graph']
                        for jitem in items_to_check:
                            if not isinstance(jitem, dict):
                                continue
                            item_type = jitem.get('@type', '')
                            if item_type in ('Product', 'IndividualProduct') or 'Product' in str(item_type):
                                product_url = jitem.get('url', '')
                                product_name = jitem.get('name', '')
                                if product_url and product_url not in seen_urls:
                                    product_url = self._make_absolute(product_url, base_url)
                                    if product_url:
                                        seen_urls.add(product_url)
                                        entries.append({
                                            "url": product_url,
                                            "title": product_name,
                                            "alt": jitem.get('description', '')[:100],
                                            "all_text": f"{product_name} {jitem.get('description', '')[:100]}",
                                        })
            except Exception as e:
                print(f"[EXTRACT] JSON-LD extraction failed: {e}")

        # Deduplicate entries with same URL (keep the one with most info)
        final = {}
        for entry in entries:
            url = entry["url"]
            if url not in final or len(entry["all_text"]) > len(final[url]["all_text"]):
                final[url] = entry

        return list(final.values())

    def _rank_entries_by_relevance(self, entries: list[dict], query: str) -> list[dict]:
        """
        Rank product entries by how well they match the query.
        Uses synonym-aware matching: "pireu" matches "piure"/"puree",
        "banane" matches "banana", etc.
        """
        query_keywords = [w.lower() for w in query.split() if len(w) >= 2]
        if not query_keywords:
            return entries

        # Pre-compute variant sets for each keyword
        kw_variant_sets = [get_word_variants(kw) for kw in query_keywords]

        def entry_score(entry: dict) -> float:
            title_lower = entry.get("title", "").lower()
            alt_lower = entry.get("alt", "").lower()
            all_text_lower = entry.get("all_text", "").lower()
            url_path = urlparse(entry["url"]).path.lower()
            url_normalized = url_path.replace("-", " ").replace("_", " ").replace("/", " ")
            combined = f"{title_lower} {alt_lower} {url_normalized} {all_text_lower}"
            entry_words = set(re.findall(r'[a-zA-ZăâîșțéèêëüöïôàáùúñçÀ-ÿ]+', combined))

            score = 0
            for kw, variants in zip(query_keywords, kw_variant_sets):
                # Exact keyword match (highest confidence)
                if kw in title_lower:
                    score += 3
                elif kw in alt_lower:
                    score += 2
                elif kw in url_normalized:
                    score += 1.5
                elif kw in all_text_lower:
                    score += 0.5
                else:
                    # Synonym/variant match (cross-language)
                    matched = variants & entry_words
                    if matched:
                        if any(v in title_lower for v in matched): score += 2.5
                        elif any(v in alt_lower for v in matched): score += 1.5
                        elif any(v in url_normalized for v in matched): score += 1.2
                        elif any(v in all_text_lower for v in matched): score += 0.4

            return score / len(query_keywords)

        return sorted(entries, key=entry_score, reverse=True)

    def _extract_product_links(self, html: str, base_url: str, query: str = "") -> list[str]:
        """Extract product page URLs from search results HTML, ranked by relevance to query."""
        entries = self._extract_product_entries(html, base_url)

        if not entries:
            return []

        if query:
            entries = self._rank_entries_by_relevance(entries, query)
            # Filter out entries with zero relevance score (no keyword match at all)
            # These are just navigation links, not actual search results
            query_keywords = [w.lower() for w in query.split() if len(w) >= 2]
            if query_keywords:
                # Pre-compute all variants for each keyword for synonym-aware matching
                kw_variant_sets = []
                for kw in query_keywords:
                    variants = get_word_variants(kw)
                    kw_variant_sets.append(variants)

                def _score(entry):
                    tl = entry.get("title", "").lower()
                    al = entry.get("alt", "").lower()
                    at = entry.get("all_text", "").lower()
                    un = urlparse(entry["url"]).path.lower().replace("-", " ").replace("_", " ")
                    combined = f"{tl} {al} {un} {at}"
                    # Tokenize combined text for word-level matching
                    entry_words = set(re.findall(r'[a-zA-ZăâîșțéèêëüöïôàáùúñçÀ-ÿ]+', combined.lower()))
                    s = 0
                    for kw, variants in zip(query_keywords, kw_variant_sets):
                        # Check exact keyword first
                        if kw in tl: s += 3
                        elif kw in al: s += 2
                        elif kw in un: s += 1.5
                        elif kw in at: s += 0.5
                        else:
                            # Check if any synonym/variant matches
                            matched_variant = variants & entry_words
                            if matched_variant:
                                # Synonym match is slightly less confident than exact
                                if any(v in tl for v in matched_variant): s += 2.5
                                elif any(v in al for v in matched_variant): s += 1.5
                                elif any(v in un for v in matched_variant): s += 1.2
                                elif any(v in at for v in matched_variant): s += 0.4
                    return s / len(query_keywords)
                entries = [e for e in entries if _score(e) > 0]
                # Sort by score descending (best match first)
                entries.sort(key=_score, reverse=True)
            # Log best match for debugging
            if entries:
                logger.debug(f"Best match for '{query}': {entries[0].get('title', '')} -> {entries[0]['url']}")

        return [e["url"] for e in entries[:10]]

    def _extract_product_images(self, product_url: str, base_url: str) -> list[str]:
        """Extract main product image(s) from a product page."""
        try:
            resp = self.session.get(product_url, timeout=6)
            if resp.status_code != 200:
                return []
        except Exception:
            return []

        images = []
        seen = set()

        def _unwrap_proxy(img_url: str) -> str:
            """Extract real image URL from proxy/CDN wrappers.
            e.g. http://proxy.com/resize?img=https://real.com/img.jpg → https://real.com/img.jpg
            """
            from urllib.parse import urlparse as _up, parse_qs
            try:
                parsed = _up(img_url)
                qs = parse_qs(parsed.query)
                # Common proxy param names: img, url, src, image, original
                for param in ('img', 'url', 'src', 'image', 'original'):
                    if param in qs:
                        candidate = qs[param][0]
                        if candidate.startswith('http') and any(
                            ext in candidate.lower() for ext in ('.jpg', '.jpeg', '.png', '.webp')
                        ):
                            return candidate
            except Exception:
                pass
            return img_url

        # Strategy 1: og:image meta tag (most reliable for main product image)
        og_match = re.findall(
            r'<meta[^>]+property=["\']og:image["\'][^>]+content=["\']([^"\']+)["\']',
            resp.text, re.IGNORECASE
        )
        og_match += re.findall(
            r'<meta[^>]+content=["\']([^"\']+)["\'][^>]+property=["\']og:image["\']',
            resp.text, re.IGNORECASE
        )
        for url in og_match:
            url = _unwrap_proxy(url)
            abs_url = self._make_absolute(url, base_url)
            if abs_url:
                abs_url = self._try_upscale_url(abs_url)
                if abs_url not in seen:
                    seen.add(abs_url)
                    images.append(abs_url)

        # Strategy 2: JSON-LD structured data (very reliable, works on many modern sites)
        json_ld_blocks = re.findall(
            r'<script[^>]*type=["\']application/ld\+json["\'][^>]*>(.*?)</script>',
            resp.text, re.IGNORECASE | re.DOTALL
        )
        for block in json_ld_blocks:
            try:
                data = json.loads(block)
                # Handle both single object and array of objects
                items = data if isinstance(data, list) else [data]
                for item in items:
                    if not isinstance(item, dict):
                        continue
                    if item.get("@type") == "Product" and item.get("image"):
                        img_val = item["image"]
                        # image can be a string, a list, or an object with url
                        img_urls = []
                        if isinstance(img_val, str):
                            img_urls = [img_val]
                        elif isinstance(img_val, list):
                            img_urls = [u for u in img_val if isinstance(u, str)]
                            img_urls += [u.get("url", "") for u in img_val if isinstance(u, dict)]
                        elif isinstance(img_val, dict):
                            img_urls = [img_val.get("url", "")]
                        for img_url in img_urls:
                            abs_url = self._make_absolute(img_url, base_url)
                            if abs_url and abs_url not in seen and self._is_product_image(abs_url):
                                # Try to get largest version by replacing size in URL
                                abs_url = self._try_upscale_url(abs_url)
                                seen.add(abs_url)
                                images.append(abs_url)
            except (json.JSONDecodeError, KeyError, TypeError):
                continue

        # Strategy 3: srcset / data-srcset — pick HIGHEST resolution variant
        # Shopify, modern sites use srcset with multiple resolutions
        srcset_matches = re.findall(
            r'(?:data-srcset|srcset)=["\']([^"\']+)["\']',
            resp.text, re.IGNORECASE
        )
        # Also extract data-zoom-image, data-large-image, data-full-image (common zoom plugins)
        hires_attrs = re.findall(
            r'<img[^>]+(?:data-zoom-image|data-large-image|data-full-image|data-zoom|data-magnify-src)=["\']([^"\']+)["\']',
            resp.text, re.IGNORECASE
        )
        for hires_url in hires_attrs:
            abs_url = self._make_absolute(hires_url, base_url)
            if abs_url and self._is_product_image(abs_url) and abs_url not in seen:
                seen.add(abs_url)
                images.insert(0, abs_url)  # High priority — insert at front

        best_srcset_url = None
        best_srcset_w = 0
        for srcset_val in srcset_matches:
            # Parse "url1 165w, url2 940w, ..." → pick largest
            candidates = []
            for part in srcset_val.split(','):
                part = part.strip()
                if not part:
                    continue
                # Format: "URL WIDTHw" or "URL Xx"
                pieces = part.rsplit(None, 1)
                if len(pieces) == 2:
                    img_url, descriptor = pieces
                    w_match = re.match(r'(\d+)w', descriptor)
                    if w_match:
                        candidates.append((int(w_match.group(1)), img_url.strip()))
                elif pieces:
                    candidates.append((0, pieces[0].strip()))
            if candidates:
                # Pick highest resolution from this srcset
                candidates.sort(key=lambda x: x[0], reverse=True)
                top_w, top_url = candidates[0]
                if top_w > best_srcset_w:
                    best_srcset_w = top_w
                    best_srcset_url = top_url

        if best_srcset_url:
            if best_srcset_url.startswith('//'):
                best_srcset_url = 'https:' + best_srcset_url
            abs_url = self._make_absolute(best_srcset_url, base_url)
            if abs_url and self._is_product_image(abs_url):
                abs_url = self._try_upscale_url(abs_url)
                if abs_url not in seen:
                    seen.add(abs_url)
                    # If srcset has larger width than og:image, put it first
                    if best_srcset_w >= 800 and images:
                        images.insert(0, abs_url)
                    else:
                        images.append(abs_url)

        # Strategy 4: Images in product media/gallery containers
        gallery_imgs = re.findall(
            r'(?:gallery|product|media|catalog)[^"]*"[^>]*>.*?<img[^>]+src=["\']([^"\']+)["\']',
            resp.text, re.IGNORECASE
        )
        # Also data-src and data-zoom-image
        gallery_imgs += re.findall(
            r'<img[^>]+(?:data-zoom-image|data-src|data-full)=["\']([^"\']+)["\']',
            resp.text, re.IGNORECASE
        )

        for url in gallery_imgs:
            abs_url = self._make_absolute(url, base_url)
            if abs_url and self._is_product_image(abs_url):
                abs_url = self._try_upscale_url(abs_url)
                if abs_url not in seen:
                    seen.add(abs_url)
                    images.append(abs_url)

        # Strategy 4: Any large image from the /media/catalog/product path (Magento)
        catalog_imgs = re.findall(
            r'["\']([^"\']*?/media/catalog/product[^"\']+)["\']',
            resp.text, re.IGNORECASE
        )
        for url in catalog_imgs:
            abs_url = self._make_absolute(url, base_url)
            if abs_url and self._is_product_image(abs_url):
                abs_url = self._try_upscale_url(abs_url)
                if abs_url not in seen:
                    seen.add(abs_url)
                    images.append(abs_url)

        # Strategy 5: Generic product image - find largest <img> on the page
        # that isn't a logo, icon, or menu image (fallback for custom platforms)
        if not images:
            all_imgs = re.findall(
                r'<img[^>]+(?:src|data-src)=["\']([^"\']+)["\'][^>]*>',
                resp.text, re.IGNORECASE
            )
            # Prefer images in /produse/, /products/, /catalog/ paths
            product_imgs = []
            other_imgs = []
            for img_tag_url in all_imgs:
                abs_url = self._make_absolute(img_tag_url, base_url)
                if not abs_url or abs_url in seen or not self._is_product_image(abs_url):
                    continue
                lower = abs_url.lower()
                if any(p in lower for p in ['/produse/', '/products/', '/product/', '/catalog/', '/uploads/product']):
                    product_imgs.append(abs_url)
                elif '/resources/' not in lower and '/header' not in lower and '/footer' not in lower:
                    other_imgs.append(abs_url)
            # Try product-path images first, then others
            for img_url in (product_imgs + other_imgs)[:3]:
                img_url = self._try_upscale_url(img_url)
                if img_url not in seen:
                    seen.add(img_url)
                    images.append(img_url)

        # Deduplicate and pick the best image (highest quality)
        # images[0] is the best candidate — could be data-zoom, srcset, og:image, or JSON-LD
        # If we have both og:image AND a higher-res version (srcset/zoom), prefer the higher-res
        if images:
            return images[:1]  # Return best single image
        return []

    def _try_upscale_url(self, url: str) -> str:
        """Try to get a larger version of an image by replacing size patterns in URL.
        - Always strips Magento cache paths (gets original full-size image).
        - For CDN thumbnails, tries larger sizes with fallback validation.
        """
        # Magento cache: /media/catalog/product/cache/<hash>/<path> → /media/catalog/product/<path>
        upscaled = re.sub(
            r'/media/catalog/product/cache/[a-f0-9]+/',
            '/media/catalog/product/',
            url
        )
        if upscaled != url:
            return upscaled

        # OpenCart/CS-Cart cache: /image/cache/...WxH.ext → try original without cache path
        oc_match = re.search(r'(/image/)cache/(.+?)-(\d+)x(\d+)\.(\w+)$', url)
        if oc_match:
            # Try original image path without cache and size suffix
            original = url[:url.index('/image/')] + oc_match.group(1) + oc_match.group(2) + '.' + oc_match.group(5)
            try:
                resp = self.session.head(original, timeout=5, allow_redirects=True)
                ct = resp.headers.get('content-type', '')
                if resp.status_code == 200 and 'image' in ct:
                    print(f"[UPSCALE] OpenCart cache → original: {original}")
                    return original
            except Exception:
                pass

        # CDN pattern: /image/WxH/ → try to upscale
        size_match = re.search(r'/image/(\d+)x(\d+)/', url)
        if size_match:
            w, h = int(size_match.group(1)), int(size_match.group(2))
            if max(w, h) < 800:
                for target in ['1080x1080', '800x800', '600x600']:
                    upscaled = re.sub(r'/image/\d+x\d+/', f'/image/{target}/', url)
                    try:
                        resp = self.session.head(upscaled, timeout=5, allow_redirects=True)
                        ct = resp.headers.get('content-type', '')
                        if resp.status_code == 200 and 'image' in ct:
                            print(f"[UPSCALE] CDN {w}x{h} → {target}: {upscaled}")
                            return upscaled
                    except Exception:
                        pass
            return url

        # Filename pattern: -WxH.ext → upscale if under 800px
        size_match = re.search(r'-(\d+)x(\d+)\.(\w+)$', url)
        if size_match:
            w, h = int(size_match.group(1)), int(size_match.group(2))
            if max(w, h) < 800:
                for target in ['1080x1080', '800x800', '600x600']:
                    upscaled = re.sub(r'-\d+x\d+\.(\w+)$', f'-{target}.\\1', url)
                    try:
                        resp = self.session.head(upscaled, timeout=5, allow_redirects=True)
                        ct = resp.headers.get('content-type', '')
                        if resp.status_code == 200 and 'image' in ct:
                            print(f"[UPSCALE] Filename {w}x{h} → {target}: {upscaled}")
                            return upscaled
                    except Exception:
                        pass
            return url

        # Shopify CDN: _WIDTHx.ext (e.g. _165x.jpg, _940x.jpg) → try larger or original
        shopify_match = re.search(r'_(\d+)x\.(\w+)(?:\?|$)', url)
        if shopify_match:
            w = int(shopify_match.group(1))
            # Always try to get the largest version — even 940x can be upscaled
            if w < 1500:
                # Try original (no size suffix), then 1500x, 1024x
                targets = []
                if w < 1500:
                    targets.append('')        # original (no size suffix)
                if w < 1500:
                    targets.append('1500x')
                if w < 1024:
                    targets.append('1024x')
                for target_w in targets:
                    if target_w:
                        upscaled = re.sub(r'_\d+x\.(\w+)', f'_{target_w}.\\1', url)
                    else:
                        # Remove size suffix entirely to get original
                        upscaled = re.sub(r'_\d+x\.(\w+)', '.\\1', url)
                    if upscaled != url:
                        try:
                            resp = self.session.head(upscaled, timeout=5, allow_redirects=True)
                            ct = resp.headers.get('content-type', '')
                            if resp.status_code == 200 and 'image' in ct:
                                print(f"[UPSCALE] Shopify {w}x → {target_w or 'original'}: {upscaled}")
                                return upscaled
                        except Exception:
                            pass
            return url

        return url

    def _make_absolute(self, url: str, base_url: str) -> str | None:
        """Convert relative URL to absolute."""
        if not url or url.startswith(('javascript:', 'mailto:', '#', 'data:')):
            return None
        if url.startswith('//'):
            return 'https:' + url
        if url.startswith('/'):
            return base_url.rstrip('/') + url
        if url.startswith('http'):
            return url
        return base_url.rstrip('/') + '/' + url

    @staticmethod
    def _domains_match(domain1: str, domain2: str) -> bool:
        """Check if two domains are the same (ignoring www. prefix and :port)."""
        d1 = domain1.lower().split(":")[0].lstrip("www.")
        d2 = domain2.lower().split(":")[0].lstrip("www.")
        return d1 == d2

    def _is_product_url(self, url: str, base_url: str) -> bool:
        """Check if URL looks like a product page (not a category, search, or static page)."""
        parsed = urlparse(url)
        base_parsed = urlparse(base_url)
        if parsed.netloc and not self._domains_match(parsed.netloc, base_parsed.netloc):
            return False
        path = parsed.path.lower()
        # Skip non-product URLs
        skip_patterns = [
            '/search', '/catalogsearch', '/cart', '/checkout', '/account',
            '/customer', '/static/', '/media/', '/js/', '/css/',
            '.js', '.css', '.png', '.jpg', '.gif', '.svg',
            '/contact', '/about', '/blog', '/category', '/wishlist',
            '/sales/', '/promo', '/quickview', '/newsletter', '/login',
            '/register', '/review', '/compare', '/faq', '/help',
            '/privacy', '/terms', '/cookie', '/sitemap',
            '/cmp/', '/page/', '/pages/', '/info/', '/cms/',
            '/brands/', '/brand/', '/producator/', '/tag/',
        ]
        if any(s in path for s in skip_patterns):
            return False
        # Must have some path depth (not just homepage)
        if len(path) <= 5:
            return False
        # Skip if path is just a single short segment (likely a category page)
        segments = [s for s in path.strip('/').split('/') if s]
        if segments and len(segments) == 1 and len(segments[0]) < 8:
            return False
        return True

    def _is_product_image(self, url: str) -> bool:
        """Check if URL looks like a product image (not icon/logo/placeholder/menu)."""
        lower = url.lower()
        skip = ['favicon', 'logo', '1x1', 'pixel', 'placeholder', '.svg', '.gif',
                'spinner', 'loading', 'blank', 'spacer', 'icon',
                '/header', '/footer', '/menu', '/banner', '/resources/images/hmm',
                '/anpc', '/gene-cms', '/wysiwyg/anpc', 'transparent.png']
        return not any(s in lower for s in skip) and any(
            ext in lower for ext in ['.jpg', '.jpeg', '.png', '.webp']
        )


# ─── IMAGE PROCESSING ────────────────────────────────────────────────────

def download_image(url: str, timeout: int = 8) -> bytes | None:
    try:
        resp = requests.get(url, timeout=timeout, headers={"User-Agent": USER_AGENT})
        resp.raise_for_status()
        ct = resp.headers.get("content-type", "")
        if "image" not in ct and "octet-stream" not in ct:
            return None
        if len(resp.content) < 2000:
            return None
        return resp.content
    except Exception:
        return None


def download_images_parallel(urls: list[str], max_workers: int = 6) -> dict[str, bytes]:
    """Download multiple images in parallel. Returns {url: data} for successful downloads."""
    from concurrent.futures import ThreadPoolExecutor, as_completed
    results = {}
    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        futures = {pool.submit(download_image, url): url for url in urls}
        for future in as_completed(futures):
            url = futures[future]
            data = future.result()
            if data:
                results[url] = data
    return results


def resize_and_pad(data: bytes, target_size: tuple[int, int],
                   bg_color: tuple = (255, 255, 255),
                   quality: int = 90,
                   remove_bg: bool = False,
                   output_format: str = "webp") -> bytes:
    """
    Process image: optional bg removal, resize with padding.
    output_format: 'webp', 'jpeg', or 'png'
    """

    if remove_bg:
        try:
            from rembg import remove as rembg_remove
            nobg = rembg_remove(data)
            img = Image.open(io.BytesIO(nobg)).convert("RGBA")
        except Exception:
            img = Image.open(io.BytesIO(data))
            img = ImageOps.exif_transpose(img)
            if img.mode != "RGBA":
                img = img.convert("RGBA")
    else:
        img = Image.open(io.BytesIO(data))
        img = ImageOps.exif_transpose(img)
        if img.mode != "RGBA":
            img = img.convert("RGBA")

    img = ImageOps.exif_transpose(img)
    original_size = max(img.size)
    tw, th = target_size

    # Use high-quality downscaling: if source is much larger, do a 2-step resize
    # (first to 2x target, then to target) for better anti-aliasing
    if original_size > max(tw, th) * 3:
        intermediate = (tw * 2, th * 2)
        img.thumbnail(intermediate, Image.LANCZOS)

    img.thumbnail((tw, th), Image.LANCZOS)

    # Apply sharpening after downscale to recover lost detail
    # UnsharpMask(radius, percent, threshold) - subtle but effective
    img = img.filter(ImageFilter.UnsharpMask(radius=1.0, percent=80, threshold=2))

    # For transparent formats (webp/png with bg removal), keep RGBA
    fmt = output_format.lower()
    if fmt == "png" and remove_bg:
        result = Image.new("RGBA", target_size, (255, 255, 255, 0))
        px = (tw - img.size[0]) // 2
        py = (th - img.size[1]) // 2
        result.paste(img, (px, py), mask=img.split()[3])
    elif fmt == "webp" and remove_bg:
        result = Image.new("RGBA", target_size, (255, 255, 255, 0))
        px = (tw - img.size[0]) // 2
        py = (th - img.size[1]) // 2
        result.paste(img, (px, py), mask=img.split()[3])
    else:
        result = Image.new("RGB", target_size, bg_color)
        px = (tw - img.size[0]) // 2
        py = (th - img.size[1]) // 2
        result.paste(img, (px, py), mask=img.split()[3])

    buf = io.BytesIO()
    if fmt == "webp":
        result.save(buf, "WEBP", quality=quality, method=4)
    elif fmt == "png":
        result.save(buf, "PNG", optimize=True)
    else:  # jpeg
        if result.mode == "RGBA":
            result = result.convert("RGB")
        result.save(buf, "JPEG", quality=quality, optimize=True)
    return buf.getvalue()


def make_thumbnail(data: bytes, size: int = 120) -> str:
    """Create a base64 thumbnail for the UI preview."""
    try:
        img = Image.open(io.BytesIO(data))
        img.thumbnail((size, size), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, "JPEG", quality=75)
        return base64.b64encode(buf.getvalue()).decode()
    except Exception:
        return ""


def url_keyword_score(url: str, product_name: str) -> float:
    """
    Score 0-1 based on how many product name keywords appear in the URL.
    Uses synonym-aware matching: "ananas" matches URL containing "pineapple".
    Only counts MEANINGFUL words (brand, flavor, type), not generic ones.
    """
    url_lower = url.lower().replace("-", " ").replace("_", " ").replace("%20", " ")
    url_words = set(re.findall(r'[a-z]+', url_lower))
    url_joined = url.lower().replace("-", "").replace("_", "").replace("%20", "")

    words = product_name.lower().split()
    # Skip short words, generic RO/EN words, and common filler
    _skip_url_words = {
        "the", "and", "for", "con", "per", "din", "cafea", "produs",
        "ceai", "infuzie", "sirop", "piure", "pireu", "aroma",
        "new", "buc", "plic", "cut", "pachet", "cutie", "set",
        "mic", "mare", "de", "cu",
    }
    keywords = [w for w in words if len(w) >= 3 and w not in _skip_url_words]

    if not keywords:
        return 0.0

    matches = 0
    for kw in keywords:
        # Direct match (exact keyword in URL)
        if kw.replace(" ", "") in url_joined:
            matches += 1
        else:
            # Synonym-aware match: check if any variant of the keyword appears
            variants = get_word_variants(kw)
            if variants & url_words:
                matches += 0.9  # Slightly less confident than exact match
    return matches / len(keywords)


def url_has_conflicting_product(url: str, product_name: str) -> bool:
    """
    Check if URL contains product-specific words (flavors, variants) that
    CONFLICT with the product name. This catches same-brand-wrong-product cases.
    e.g. searching "Tea Tales Blueberry Cream" but URL has "green-jasmine"

    Returns True if a conflict is detected (image is likely wrong product).
    """
    # Include parenthetical text as it often contains the flavor/variant
    # e.g. "Pireu MONIN Fructe de Padure (Red Berries) 1L" — "Red Berries" IS the identity
    url_lower = url.lower().replace("-", " ").replace("_", " ").replace("%20", " ")
    url_words = set(re.findall(r'[a-z]{3,}', url_lower))
    name_words = set(re.findall(r'[a-z]{3,}', product_name.lower()))

    # PRODUCT TYPE CONFLICT: detect sirop vs piure, ceai vs cafea, etc.
    # If name says "pireu/puree" but URL has "sirop/syrup", it's the wrong product type.
    _type_groups = [
        {"pireu", "piure", "puree", "fructe"},       # puree group
        {"sirop", "syrup", "sirops"},                 # syrup group
        {"ceai", "tea", "infuzie", "infusion"},       # tea group
        {"cafea", "coffee"},                          # coffee group
        {"sos", "sauce"},                             # sauce group
        {"gem", "jam", "dulceata"},                   # jam group
        {"suc", "juice"},                             # juice group
    ]
    name_type_group = None
    for group in _type_groups:
        if name_words & group:
            name_type_group = group
            break
    if name_type_group:
        # Check if URL contains words from a DIFFERENT type group
        # BUT only flag conflict if the URL does NOT also contain the correct type
        # (page titles may include navigation breadcrumbs with multiple type words)
        url_has_correct_type = url_words & name_type_group
        for group in _type_groups:
            if group is name_type_group:
                continue
            url_has_other_type = url_words & group
            if url_has_other_type and not url_has_correct_type:
                print(f"[CONFLICT] Type mismatch: name type={name_type_group & name_words}, URL has={url_has_other_type}")
                return True

    # Common flavor/variant words that distinguish products within the same brand
    # If any of these appear in URL but NOT in product name, it's a different product
    _flavor_words = {
        # Fruits
        "mango", "banana", "strawberry", "blueberry", "raspberry", "cherry",
        "apple", "peach", "lemon", "lime", "orange", "grapefruit", "pineapple",
        "coconut", "passion", "kiwi", "melon", "watermelon", "pomegranate",
        "cranberry", "blackberry", "apricot", "guava", "papaya", "tangerine",
        "mandarin", "fig", "plum", "grape", "pear", "litchi", "lychee",
        # RO fruits
        "capsuni", "capsuna", "afine", "afina", "cirese", "ciresa",
        "piersici", "piersica", "lamaie", "lamaia", "portocale", "portocala",
        "banane", "mandarine", "mandarina", "mere", "ananas", "zmeura",
        "mure", "caise", "coacaze",
        # Flavors/variants
        "vanilla", "vanilie", "chocolate", "ciocolata", "caramel",
        "cinnamon", "scortisoara", "ginger", "mint", "menta",
        "jasmine", "iasomie", "chamomile", "musetel", "lavender", "lavanda",
        "rose", "trandafir", "hibiscus", "honey", "miere",
        "hazelnut", "alune", "almond", "migdale", "pistachio", "fistic",
        "cream", "butterscotch", "toffee", "cookie",
        # Colors that indicate flavor
        "green", "red", "blue", "black", "white",
    }

    # Find flavor words in URL that are NOT in product name
    url_flavors = url_words & _flavor_words
    name_flavors = name_words & _flavor_words

    # Also check variant translations
    name_flavor_variants = set()
    for nf in name_flavors:
        name_flavor_variants |= get_word_variants(nf)

    # Conflict: URL has a flavor word that's NOT any variant of product name flavors
    conflicting = url_flavors - name_flavor_variants - name_words
    if conflicting:
        # Make sure the conflicting word isn't just part of a brand name
        # (e.g. "green" in "Green Mountain Coffee" isn't a flavor conflict)
        # Only flag if product name ALSO has distinct flavor words
        if name_flavors:
            return True

    # PRODUCT LINE CHECK: if product name has distinctive multi-word identity
    # (e.g. "PAULISTA" in "LAVAZZA PAULISTA"), check if URL has a DIFFERENT
    # product line name instead (e.g. "crema-gusto" instead of "paulista")
    _generic_words = {
        "cafea", "coffee", "ceai", "tea", "boabe", "beans", "macinata", "ground",
        "capsule", "capsules", "instant", "solubila", "plic", "pliculete",
        "sirop", "syrup", "piure", "pireu", "puree", "cutie", "pachet",
        "infuzie", "infusion", "looseleaf", "loose", "leaf", "herbal",
        "frunze", "leaves",
    }
    _size_words = {"mic", "mare", "mediu", "mini", "nou", "new"}
    # Descriptor/modifier words — not distinctive enough to identify a product
    _descriptor_words = {
        "negru", "alb", "rosu", "verde", "galben",  # RO colors
        "premium", "clasic", "original", "traditional", "special",
        "organic", "bio", "eco", "natural", "artisan",
    }
    # URL infrastructure words — appear in image/CDN paths, NOT product identifiers
    _infra_words = {
        "media", "cache", "catalog", "product", "image", "images", "content",
        "uploads", "files", "storage", "thumb", "thumbnail", "resize",
        "static", "assets", "dist", "public", "data", "default", "large",
        "small", "medium", "category", "jpeg", "webp", "https", "http", "html",
        "store", "root", "original", "placeholder", "watermark", "scaled",
        "photos", "produse", "domains", "coreimg", "cust", "single",
        "solution", "business", "carousel", "customcontainer", "boardimages",
        "pret", "price", "produs", "promo", "reducere", "stoc", "livrare",
    }
    _ro_en_values = set(v.lower() for v in _RO_EN_MAP.values())
    _all_ignore = _generic_words | _size_words | _flavor_words | _ro_en_values | _infra_words | _descriptor_words

    # Use only the URL PATH (strip domain, scheme, query) to avoid false positives
    try:
        from urllib.parse import urlparse as _uparse
        url_path_text = _uparse(url_lower.replace(" ", "-")).path.replace("-", " ").replace("_", " ").replace("/", " ")
    except Exception:
        url_path_text = url_lower
    url_path_words = set(re.findall(r'[a-z]{4,}', url_path_text))

    distinctive_name_words = set()
    for w in name_words:
        if len(w) >= 4 and w not in _all_ignore and not re.match(r'^\d', w):
            distinctive_name_words.add(w)

    if distinctive_name_words and len(distinctive_name_words) >= 1:
        url_path_distinctive = set()
        for w in url_path_words:
            if len(w) >= 4 and w not in _all_ignore and not re.match(r'^\d', w):
                url_path_distinctive.add(w)

        # Check name words against URL using SUBSTRING matching
        # Also check RO↔EN translations (e.g. "pere" should match "pear" in URL)
        name_found_in_url = set()
        for nw in distinctive_name_words:
            if nw in url_path_text:
                name_found_in_url.add(nw)
            else:
                # Try translation: RO→EN or EN→RO
                en_translation = _RO_EN_MAP.get(nw)
                if en_translation and en_translation.lower() in url_path_text:
                    name_found_in_url.add(nw)
                else:
                    # Try all variants (includes multi-language synonyms)
                    for variant in get_word_variants(nw):
                        if variant in url_path_text:
                            name_found_in_url.add(nw)
                            break
        name_missing = distinctive_name_words - name_found_in_url

        # URL distinctive words not in product name (even as substring)
        # Also check if URL word is a CONCATENATION of consecutive name words
        # e.g. "teatales" = "tea" + "tales", "droetker" = "dr" + "oetker"
        name_words_list = re.findall(r'[a-z]+', product_name.lower())
        url_only = set()
        name_lower_joined = product_name.lower()
        for uw in url_path_distinctive:
            if uw in name_lower_joined:
                continue
            # Check if uw is a concatenation of 2-3 consecutive name words
            is_concat = False
            for i in range(len(name_words_list)):
                concat = ""
                for j in range(i, min(i + 4, len(name_words_list))):
                    concat += name_words_list[j]
                    if concat == uw:
                        is_concat = True
                        break
                if is_concat:
                    break
            if not is_concat:
                url_only.add(uw)

        # Conflict if: URL has foreign words AND name has missing words
        # OR: name has 2+ distinctive words missing from URL (even if no foreign URL words)
        # The second case catches brand-match-but-wrong-product: same brand, different product line
        if url_only and name_missing and len(name_missing) >= 1:
            print(f"[CONFLICT] Product line mismatch: name has {name_missing} (missing from URL), URL has {url_only} (not in name)")
            return True
        if not url_only and name_missing and len(name_missing) >= 2:
            # Many distinctive name words missing from URL — likely wrong product
            # even though URL doesn't have "foreign" words
            print(f"[CONFLICT] Product identity mismatch: {name_missing} missing from URL (no foreign words but too many missing)")
            return True

    return False


# ─── SMART QUERY BUILDER ─────────────────────────────────────────────

# Packaging noise patterns to remove from search queries
_NOISE_PATTERNS = [
    r'\bNEW\b',                          # "NEW" marker
    r'\(\d+\s*plic[/\\]cut\)',           # (20plic/cut)
    r'\(\d+\s*buc[/\\]?[a-z]*\)',        # (10buc/pac)
    r'\(\d+\s*[a-z]*\)',                 # (20capsule) etc.
    r'\b\d+[.,]\d+\s*(gr|g|kg|ml|l|cl)\b',  # 0.33l, 1.5kg (MUST be before simpler pattern)
    r'\b\d+x\d+\s*(gr|g|kg|ml|l|cl)?\b',    # 6x330ml
    r'\b\d+\s*(gr|g|kg|ml|l|cl)\b',     # 4gr, 250g, 0.75l, 33cl
    r'\b\d+[.,]?\d*\s*$',               # Trailing numbers/decimals left over
    r'\b(mic|mare|mediu|mini)\b',        # Size words: mic=small, mare=large
]

# Multi-language product translations: RO → EN (primary), plus FR/DE/ES/IT synonyms
# Used for: cross-language search queries and CLIP prompt building
# Each key (RO) maps to EN. _LANG_SYNONYMS maps EN → all language variants for matching.
_RO_EN_MAP = {
    "ceai": "tea",
    "cafea": "coffee",
    "ciocolata": "chocolate",
    "biscuiti": "biscuits",
    "napolitane": "wafers",
    "suc": "juice",
    "apa": "water",
    "bere": "beer",
    "vin": "wine",
    "lapte": "milk",
    "ulei": "oil",
    "otet": "vinegar",
    "zahar": "sugar",
    "faina": "flour",
    "orez": "rice",
    "paste": "pasta",
    "sos": "sauce",
    "conserva": "canned",
    "paine": "bread",
    "negru": "black",
    "verde": "green",
    "infuzie": "infusion",
    "sirop": "syrup",
    "piure": "puree",
    "prajitura": "cake",
    "frisca": "cream",
    "smantana": "cream",
    "unt": "butter",
    "miere": "honey",
    "gem": "jam",
    "compot": "compote",
    "inghetata": "ice cream",
    "alune": "hazelnuts",
    "migdale": "almonds",
    "nuci": "walnuts",
    "capsuni": "strawberry",
    "zmeura": "raspberry",
    "afine": "blueberry",
    "visine": "sour cherry",
    "cirese": "cherry",
    "portocale": "orange",
    "lamaie": "lemon",
    "mar": "apple",
    "pere": "pear",
    "piersici": "peach",
    "banane": "banana",
    "ananas": "pineapple",
    "mango": "mango",
    "cocos": "coconut",
    "vanilie": "vanilla",
    "scortisoara": "cinnamon",
    "ghimbir": "ginger",
    "menta": "mint",
    "caramel": "caramel",
    "cacao": "cocoa",
    "pireu": "puree",
    "mere": "apple",
    "rosii": "tomato",
    "castraveti": "cucumber",
    "morcovi": "carrot",
    "mazare": "peas",
    "fasole": "beans",
    "ardei": "pepper",
    "usturoi": "garlic",
    "ceapa": "onion",
    "ciuperci": "mushroom",
    "padure": "forest",
    "fructe": "fruit",
    "mandarine": "tangerine",
    "mandarina": "tangerine",
    "mandarin": "tangerine",
    "grepfrut": "grapefruit",
    "kiwi": "kiwi",
    "pepene": "watermelon",
    "caise": "apricot",
    "prune": "plum",
    "smochine": "fig",
    "rodii": "pomegranate",
    "trandafir": "rose",
    "lavanda": "lavender",
    "fructul": "fruit",
    "pasiunii": "passion",
    "passion": "passion",
    "zmeura": "raspberry",
}

# Synonyms across languages: EN keyword → [FR, DE, ES, IT, RO variants]
# Used to generate extra search query variants
_LANG_SYNONYMS = {
    "syrup":      ["sirop", "sirup", "sirope", "sciroppo"],
    "puree":      ["piure", "püree", "puré", "purée"],
    "tea":        ["ceai", "tee", "té", "tè", "thé"],
    "coffee":     ["cafea", "kaffee", "café"],
    "chocolate":  ["ciocolata", "schokolade", "chocolat"],
    "juice":      ["suc", "saft", "jus", "jugo", "succo"],
    "cream":      ["crème", "creme", "crema", "frisca", "smantana"],
    "butter":     ["unt", "beurre", "mantequilla", "burro"],
    "honey":      ["miere", "miel", "honig"],
    "jam":        ["gem", "confiture", "marmelade", "mermelada", "marmellata"],
    "compote":    ["compot", "kompott"],
    "cake":       ["prajitura", "gâteau", "kuchen", "torta", "pastel"],
    "vanilla":    ["vanilie", "vanille", "vainilla", "vaniglia"],
    "cinnamon":   ["scortisoara", "cannelle", "zimt", "canela", "cannella"],
    "ginger":     ["ghimbir", "gingembre", "ingwer", "jengibre", "zenzero"],
    "mint":       ["menta", "menthe", "minze", "menta", "hierbabuena"],
    "strawberry": ["capsuni", "capsuna", "capsune", "fraise", "erdbeere", "fresa", "fragola"],
    "raspberry":  ["zmeura", "framboise", "himbeere", "frambuesa", "lampone"],
    "blueberry":  ["afine", "afina", "myrtille", "blaubeere", "arándano", "mirtillo"],
    "cherry":     ["cirese", "ciresa", "cerise", "kirsche", "cereza", "ciliegia"],
    "lemon":      ["lamaie", "lamaia", "lamai", "citron", "zitrone", "limón", "limone"],
    "orange":     ["portocale", "portocala", "orange", "naranja", "arancia"],
    "banana":     ["banane", "banana", "banane", "plátano"],
    "peach":      ["piersici", "piersica", "pêche", "pfirsich", "melocotón", "pesca"],
    "apple":      ["mar", "mere", "pomme", "apfel", "manzana", "mela"],
    "tangerine":  ["mandarine", "mandarina", "mandarin", "clémentine", "clementina"],
    "coconut":    ["cocos", "coco", "kokosnuss"],
    "caramel":    ["caramel", "karamell", "caramelo", "caramello"],
    "hazelnut":   ["alune", "noisette", "haselnuss", "avellana", "nocciola"],
    "almond":     ["migdale", "amande", "mandel", "almendra", "mandorla"],
    "walnut":     ["nuci", "noix", "walnuss", "nuez", "noce"],
    "infusion":   ["infuzie", "infusion", "aufguss", "infusión", "infusione"],
    "pineapple":  ["ananas"],
    "mango":      ["mango"],
    "passion fruit": ["fructul pasiunii", "maracuja", "passion"],
    "watermelon": ["pepene"],
    "apricot":    ["caise"],
    "plum":       ["prune"],
    "fig":        ["smochine"],
    "pomegranate":["rodii", "rodie"],
    "rose":       ["trandafir"],
    "lavender":   ["lavanda"],
    "fruit":      ["fructe"],
}

# ── WORD VARIANT SYSTEM ──────────────────────────────────────────────────
# Maps any word (RO, EN, FR, etc.) to ALL equivalent forms.
# e.g. "pireu" → {"pireu", "piure", "puree", "purée", "püree", "puré"}
# e.g. "banane" → {"banane", "banana", "plátano"}

_WORD_VARIANTS_CACHE: dict[str, set[str]] = {}

def _build_word_variants_map():
    """Build a lookup: lowercase word → set of all equivalent forms."""
    if _WORD_VARIANTS_CACHE:
        return  # Already built

    # Group all equivalent words together
    groups: list[set[str]] = []

    for en_key, synonyms in _LANG_SYNONYMS.items():
        group = {en_key.lower()}
        for s in synonyms:
            group.add(s.lower())
        # Add RO variants from _RO_EN_MAP that map to this EN key
        for ro, en in _RO_EN_MAP.items():
            if en.lower() == en_key.lower():
                group.add(ro.lower())
        groups.append(group)

    # Add RO→EN pairs not covered by _LANG_SYNONYMS
    covered_en = {en_key.lower() for en_key in _LANG_SYNONYMS}
    for ro, en in _RO_EN_MAP.items():
        if en.lower() not in covered_en:
            groups.append({ro.lower(), en.lower()})

    # Build lookup: word → group
    for group in groups:
        frozen = frozenset(group)
        for word in group:
            if word in _WORD_VARIANTS_CACHE:
                # Merge groups
                existing = _WORD_VARIANTS_CACHE[word]
                merged = existing | group
                for w in merged:
                    _WORD_VARIANTS_CACHE[w] = merged
            else:
                _WORD_VARIANTS_CACHE[word] = set(group)


def get_word_variants(word: str) -> set[str]:
    """Get all equivalent forms of a word across languages.
    e.g. get_word_variants("pireu") → {"pireu", "piure", "puree", "purée", ...}
    """
    _build_word_variants_map()
    return _WORD_VARIANTS_CACHE.get(word.lower(), {word.lower()})


def words_match(word1: str, word2: str) -> bool:
    """Check if two words are equivalent (same meaning in any language).
    e.g. words_match("pireu", "piure") → True
         words_match("banane", "banana") → True
    """
    if word1.lower() == word2.lower():
        return True
    variants = get_word_variants(word1)
    return word2.lower() in variants


def normalize_to_english(word: str) -> str | None:
    """Get the English equivalent of a word, or None if not found.
    e.g. normalize_to_english("pireu") → "puree"
         normalize_to_english("banane") → "banana"
    """
    lower = word.lower()
    # Direct RO→EN lookup
    en = _RO_EN_MAP.get(lower)
    if en:
        return en
    # Check if it's already an EN key in _LANG_SYNONYMS
    if lower in _LANG_SYNONYMS:
        return lower
    # Reverse lookup in _LANG_SYNONYMS
    for en_key, synonyms in _LANG_SYNONYMS.items():
        if lower in [s.lower() for s in synonyms]:
            return en_key
    return None


def build_direct_search_queries(cleaned_name: str, key_words: list[str]) -> list[str]:
    """
    Build smart search query variants for direct site scraping.
    Key insight: site search engines work best with SIMPLE queries (brand + flavor).
    Product type words (piure, sirop, ceai) often CONFUSE site search.

    For "Pireu MONIN Banane" we generate:
      1. "MONIN banana" (brand + EN flavor — BEST for most sites)
      2. "MONIN Banane" (brand + original flavor)
      3. "MONIN" (brand only — broad fallback)
      4. "Pireu MONIN Banane" (full name — only if specific search works)
    """
    queries = []
    words = cleaned_name.split()
    filler_words = {"de", "cu", "si", "din", "la", "pt", "pentru", "si", "sau", "mic", "mare", "mediu", "mini"}

    # Product type words: these describe WHAT it is, not WHICH one
    # They confuse most site search engines, so we strip them for primary queries
    product_type_en = {"puree", "syrup", "tea", "coffee", "infusion", "juice",
                       "sauce", "cream", "jam", "compote", "butter", "honey",
                       "chocolate", "cake", "ice cream"}
    product_type_ro = set()
    for ro, en in _RO_EN_MAP.items():
        if en.lower() in product_type_en:
            product_type_ro.add(ro)

    # Classify words into: brand, flavor, product_type
    brand_words = []
    flavor_words = []  # (original, en_translation)
    type_words = []    # (original, en_translation)
    for w in words:
        lower = w.lower()
        if lower in filler_words or len(w) <= 1:
            continue
        if re.match(r'^[\d.]+$', w):
            continue
        en = normalize_to_english(lower)
        if en and en.lower() != lower:
            # Translatable word — is it a product type or a flavor?
            if en.lower() in product_type_en or lower in product_type_ro:
                type_words.append((w, en))
            else:
                flavor_words.append((w, en))
        else:
            brand_words.append(w)

    # === QUERY STRATEGIES (ordered by search-engine friendliness) ===

    # Strategy 1: Brand + EN flavor only → "MONIN banana" (BEST for most sites)
    if brand_words and flavor_words:
        en_flavors = " ".join(fw[1] for fw in flavor_words)
        queries.append(f"{' '.join(brand_words)} {en_flavors}")

    # Strategy 2: Brand + original flavor → "MONIN Banane"
    if brand_words and flavor_words:
        orig_flavors = " ".join(fw[0] for fw in flavor_words)
        queries.append(f"{' '.join(brand_words)} {orig_flavors}")

    # Strategy 3: Just brand → "MONIN" (very broad, needs AI matching to pick right product)
    if brand_words:
        queries.append(" ".join(brand_words))

    # Strategy 4: Concatenated brand + flavor → "TeaTales Blueberry Cream"
    if len(brand_words) >= 2:
        merged = brand_words[0] + brand_words[1]
        rest_brand = brand_words[2:]
        brand_str = f"{merged} {' '.join(rest_brand)}" if rest_brand else merged
        if flavor_words:
            en_flavors = " ".join(fw[1] for fw in flavor_words)
            queries.append(f"{brand_str} {en_flavors}")
        else:
            queries.append(brand_str)

    # Strategy 5: Brand + type + flavor (EN) → "MONIN puree banana"
    if brand_words and type_words and flavor_words:
        en_type = " ".join(tw[1] for tw in type_words)
        en_flavors = " ".join(fw[1] for fw in flavor_words)
        queries.append(f"{' '.join(brand_words)} {en_type} {en_flavors}")

    # Strategy 6: Full cleaned name → "Pireu MONIN Banane"
    queries.append(cleaned_name)

    # Strategy 7: All-EN translation → "puree MONIN banana"
    all_en = []
    has_any_translation = False
    for w in words:
        lower = w.lower()
        if lower in filler_words:
            continue
        en = normalize_to_english(lower)
        if en and en.lower() != lower:
            all_en.append(en)
            has_any_translation = True
        else:
            all_en.append(w)
    if has_any_translation:
        queries.append(" ".join(all_en))

    # Strategy 8: RO synonym variants (e.g. "pireu" → try "piure")
    for fw_orig, fw_en in flavor_words:
        variants = get_word_variants(fw_orig.lower())
        for variant in variants:
            if variant != fw_orig.lower() and variant != fw_en.lower():
                if brand_words:
                    queries.append(f"{' '.join(brand_words)} {variant}")
                break

    # Deduplicate while preserving order, cap at 3 queries for speed
    seen = set()
    unique = []
    for q in queries:
        q_clean = re.sub(r'\s+', ' ', q).strip()
        q_key = q_clean.lower()
        if q_key and q_key not in seen:
            seen.add(q_key)
            unique.append(q_clean)
    return unique[:5]  # Top 5 queries; more is diminishing returns


def clean_product_query(denumire: str) -> str:
    """
    Remove packaging noise from product name for better search.
    'Ceai Tea Tales 4gr Blueberry Cream NEW (20plic/cut)' → 'Ceai Tea Tales Blueberry Cream'
    """
    cleaned = denumire
    # Remove parenthetical hints like "(mere verzi)", "(20plic/cut)"
    cleaned = re.sub(r'\([^)]*\)', '', cleaned)
    for pattern in _NOISE_PATTERNS:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)
    # Remove extra whitespace
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned


def build_search_variants(denumire: str, search_suffix: str) -> list[str]:
    """
    Generate multiple search query variants for better coverage:
    1. Cleaned exact name (quoted)
    2. Cleaned name + suffix
    3. Brand + flavor only (most important keywords)
    4. Cross-language variant (RO words → EN or vice-versa)
    """
    cleaned = clean_product_query(denumire)
    words = cleaned.split()
    variants = []

    # 1. Cleaned exact (quoted)
    variants.append(f'"{cleaned}"')

    # 2. Cleaned + suffix
    variants.append(f"{cleaned} {search_suffix}".strip())

    # 3. Brand + flavor: skip generic RO words, keep brand/flavor
    skip_words = set(_RO_EN_MAP.keys()) | {"de", "cu", "si", "din", "la", "pt", "pentru"}
    key_words = [w for w in words if w.lower() not in skip_words and len(w) > 1]
    if key_words and len(key_words) < len(words):
        brand_flavor = " ".join(key_words)
        variants.append(f"{brand_flavor} {search_suffix}".strip())

    # 4. Cross-language: translate RO product type words to EN
    # Skip translation if the English word already exists in the name
    all_words_lower = {w.lower() for w in words}
    translated_words = []
    has_translation = False
    for w in words:
        lower_w = w.lower()
        en_word = _RO_EN_MAP.get(lower_w)
        if en_word and en_word.lower() not in all_words_lower:
            translated_words.append(en_word)
            has_translation = True
        elif en_word and en_word.lower() in all_words_lower:
            # RO word is redundant (EN equivalent already present), skip it
            has_translation = True
        else:
            translated_words.append(w)
    if has_translation:
        cross_lang = " ".join(translated_words)
        variants.append(f"{cross_lang} {search_suffix}".strip())

    # 5. Short: first meaningful word (brand) + last meaningful words (flavor)
    if len(key_words) > 3:
        short = f"{key_words[0]} {' '.join(key_words[-2:])}"
        variants.append(f"{short} {search_suffix}".strip())

    # 6. Concatenated brand variant: "Tea Tales" → "TeaTales" / "TEATALES"
    if len(key_words) >= 2:
        merged_brand = key_words[0] + key_words[1]
        rest_words = key_words[2:]
        if rest_words:
            concat_variant = f"{merged_brand} {' '.join(rest_words)}"
        else:
            concat_variant = merged_brand
        variants.append(concat_variant)

    # 7. Multi-language variants (FR, DE, ES, IT)
    # e.g. "piure" → search also with "purée", "püree"
    # Build a reverse lookup: any language → EN
    _all_to_en = {}
    for en_key, synonyms in _LANG_SYNONYMS.items():
        _all_to_en[en_key] = en_key
        for syn in synonyms:
            _all_to_en[syn.lower()] = en_key
    for ro_key, en_key in _RO_EN_MAP.items():
        _all_to_en[ro_key] = en_key

    for w in words:
        lower_w = w.lower()
        en_key = _all_to_en.get(lower_w)
        if en_key and en_key in _LANG_SYNONYMS:
            # Generate a variant for each language synonym
            for synonym in _LANG_SYNONYMS[en_key]:
                if synonym.lower() != lower_w:
                    # Replace this word with the synonym
                    lang_words = [synonym if ww.lower() == lower_w else ww for ww in words]
                    lang_variant = " ".join(lang_words)
                    variants.append(f"{lang_variant} {search_suffix}".strip())
                    break  # Just add the first alternative (usually FR, most common)

    # Deduplicate while keeping order
    seen = set()
    unique = []
    for v in variants:
        if v not in seen:
            seen.add(v)
            unique.append(v)
    return unique


def sanitize_comment(denumire: str) -> str:
    s = denumire.lower().strip()
    s = (s.replace(' ', '').replace('ă', 'a').replace('â', 'a')
         .replace('î', 'i').replace('ș', 's').replace('ş', 's')
         .replace('ț', 't').replace('ţ', 't'))
    s = re.sub(r'[^a-z0-9._-]', '', s)
    return s[:80]


def hermes_filename(product_id: str, denumire: str, image_number: int,
                    fmt: str = "webp") -> str:
    comment = sanitize_comment(denumire)
    ext = {"webp": ".webp", "png": ".png", "jpeg": ".jpg"}.get(fmt.lower(), ".webp")
    return str(product_id) + "}{" + comment + "}#" + str(image_number) + ext


# ─── SCRAPER JOB ──────────────────────────────────────────────────────────

def run_scraper_job(job_id: str, products: list[dict], config: dict,
                    event_queue: queue.Queue):
    """Runs the full scraping pipeline in a background thread."""

    output_dir = OUTPUT_DIR / job_id
    output_dir.mkdir(parents=True, exist_ok=True)

    # Config
    target_size = (config.get("image_width", 200), config.get("image_height", 200))
    quality = config.get("quality", 98)
    remove_bg = config.get("remove_background", False)
    output_format = config.get("output_format", "png")
    images_per_product = config.get("images_per_product", 1)
    search_suffix = config.get("search_suffix", "product photo")
    pexels_key = config.get("pexels_key", "")
    bing_key = config.get("bing_key", "")
    max_candidates = config.get("max_candidates", 10)
    # Clean priority sites: extract domain from full URLs
    raw_sites = [s.strip() for s in config.get("priority_sites", []) if s.strip()]
    priority_sites = []
    for site in raw_sites:
        # Remove protocol and path, keep only domain
        s = site.replace("https://", "").replace("http://", "")
        s = s.split("/")[0]  # Remove path
        s = s.split("?")[0]  # Remove query params
        s = s.lstrip("www.")  # Optional: remove www
        if s:
            priority_sites.append(s)

    quality_checker = ImageQualityChecker({
        "min_resolution": config.get("min_resolution", 200),
        "min_quality_score": config.get("min_quality_score", 40),
        "reject_blurry": config.get("reject_blurry", True),
        "min_aspect_ratio": config.get("min_aspect_ratio", 0.4),
        "max_aspect_ratio": config.get("max_aspect_ratio", 2.5),
    })

    def send(event_type, data):
        event_queue.put({"event": event_type, "data": data})

    # Relevance checker (CLIP)
    use_relevance = config.get("check_relevance", True)
    min_relevance = config.get("min_relevance", 0.20)
    relevance_checker = get_relevance_checker(min_relevance) if use_relevance else None

    # Pre-load CLIP model before starting (so first product isn't slow)
    if relevance_checker:
        send("status", {"message": "Loading AI relevance model (first time only)..."})
        relevance_checker._load()
        if relevance_checker.available:
            send("status", {"message": "AI relevance model ready"})
        else:
            send("status", {"message": "AI model not available, skipping relevance check"})
            relevance_checker = None

    # Init search engines
    bing_scraper = BingImageScraper()
    ai_matcher = AIProductMatcher(config.get("anthropic_key", ""))
    direct_scraper = DirectSiteScraper(ai_matcher=ai_matcher)
    ddg = DuckDuckGoSearch()
    pexels = PexelsSearch(pexels_key)

    if ai_matcher.available:
        send("status", {"message": "AI product matching enabled (Claude Haiku)"})

    gemini_key = config.get("gemini_key", "") or os.environ.get("GEMINI_API_KEY", "") or _load_config_file().get("gemini_key", "")
    print(f"[CONFIG] gemini_key present: {bool(gemini_key)} (len={len(gemini_key) if gemini_key else 0})")
    print(f"[CONFIG] config keys: {list(config.keys())}")
    if gemini_key:
        send("status", {"message": "Gemini Vision image verification enabled (FREE)"})
    else:
        send("status", {"message": "Gemini Vision disabled — add Gemini API key for image verification (free at aistudio.google.com)"})

    send("job_start", {
        "job_id": job_id,
        "total_products": len(products),
        "config": config,
    })

    results = []
    stats = {"total": len(products), "success": 0, "failed": 0, "images_saved": 0}

    for idx, product in enumerate(products):
        # Check if job was cancelled
        job = active_jobs.get(job_id)
        if job and job.get("status") == "cancelled":
            send("job_done", {
                "job_id": job_id, "stats": stats,
                "output_dir": str(output_dir.resolve()),
                "cancelled": True,
            })
            return

        denumire = product["denumire"].strip()
        product_id = product.get("id", str(idx + 1))

        if not denumire:
            continue

        send("product_start", {
            "index": idx,
            "total": len(products),
            "product_id": product_id,
            "denumire": denumire,
        })

        # Build search queries - smart cleaning + cross-language variants
        GREAT_RELEVANCE = 80

        query_variants = build_search_variants(denumire, search_suffix)

        # ── SEARCH: collect URLs from multiple queries ──
        search_urls = []  # [(url, source)]
        seen_urls = set()

        img_to_product_url = {}  # image_url → product_page_url (for conflict checking)

        def collect(results, source):
            for r in results:
                url = r.get("image", "")
                if url and url not in seen_urls:
                    seen_urls.add(url)
                    search_urls.append((url, source))
                    # Store product_url mapping for later conflict checking
                    if r.get("product_url"):
                        img_to_product_url[url] = r["product_url"]

        # Priority sites first - direct scraping (ALL SITES IN PARALLEL)
        if priority_sites:
            cleaned_name = clean_product_query(denumire)
            filler_words = {"de", "cu", "si", "din", "la", "pt", "pentru", "si", "sau", "mic", "mare", "mediu", "mini"}
            skip_words = set(_RO_EN_MAP.keys()) | filler_words
            key_words = [w for w in cleaned_name.split() if w.lower() not in skip_words and len(w) > 1]
            direct_queries = build_direct_search_queries(cleaned_name, key_words)

            def _search_one_site(site):
                """Search a single priority site. Returns (site, results_list).
                FAST: if first query finds results but ALL are wrong type/conflict,
                skip remaining queries for this site."""
                site_had_results = False
                for dq in direct_queries:
                    try:
                        sr = direct_scraper.search(site, dq, max_results=max_candidates,
                                                   product_name=denumire)
                        if sr:
                            site_had_results = True
                            # Final safety: filter out conflicting products
                            filtered_sr = []
                            for r in sr:
                                prod_url = r.get('product_url', '')
                                if prod_url and url_has_conflicting_product(prod_url, denumire):
                                    continue
                                title = r.get('title', '')
                                if title and url_has_conflicting_product(title, denumire):
                                    continue
                                filtered_sr.append(r)
                            if filtered_sr:
                                for r in filtered_sr:
                                    print(f"[PRIORITY] {site} query='{dq}' → image={r.get('image','')[:80]} | product_url={r.get('product_url','')[:80]}")
                                return (site, filtered_sr)
                            else:
                                # Site has products but ALL conflict — don't try more queries
                                print(f"[PRIORITY] {site} query='{dq}' → all {len(sr)} results conflicted, skipping site")
                                return (site, [])
                    except Exception as e:
                        logger.debug(f"Direct scrape failed for {site}: {e}")
                return (site, [])

            # Search ALL priority sites in parallel
            from concurrent.futures import ThreadPoolExecutor, as_completed
            send("search_phase", {"index": idx, "phase": "priority", "site": ", ".join(priority_sites), "query": "parallel direct search"})
            with ThreadPoolExecutor(max_workers=min(len(priority_sites), 5)) as pool:
                futures = {pool.submit(_search_one_site, site): site for site in priority_sites}
                for future in as_completed(futures):
                    site, sr = future.result()
                    if sr:
                        collect(sr, f"direct:{site}")

            # === PHASE 2: Search engine fallback (if NO site returned results) ===
            if not search_urls:
                for site in priority_sites[:2]:  # Limit fallback to first 2 sites
                    site_queries = [
                        f"site:{site} {cleaned_name}",
                    ]
                    if key_words and len(key_words) < len(cleaned_name.split()):
                        site_queries.append(f"site:{site} {' '.join(key_words)}")
                    if len(key_words) >= 2:
                        merged = key_words[0] + key_words[1]
                        rest = key_words[2:]
                        concat_q = f"{merged} {' '.join(rest)}" if rest else merged
                        site_queries.append(f"site:{site} {concat_q}")

                    for sq in dict.fromkeys(site_queries):
                        send("search_phase", {"index": idx, "phase": "priority", "site": site, "query": sq})
                        for sn, searcher in [("bing", bing_scraper), ("ddg", ddg)]:
                            try:
                                sr = searcher.search(sq, max_results=max_candidates)
                                if sr:
                                    collect(sr, f"{sn}:{site}")
                                    break
                            except Exception:
                                pass
                        if len(search_urls) >= max_candidates:
                            break

        # General search with multiple query variants
        for qi, query in enumerate(query_variants):
            if len(search_urls) >= max_candidates:
                break
            if qi == 0:
                send("search_phase", {"index": idx, "phase": "general", "query": query})
            for sn, searcher in [("bing", bing_scraper), ("ddg", ddg), ("pexels", pexels)]:
                if sn == "pexels" and not pexels.available:
                    continue
                try:
                    sr = searcher.search(query, max_results=max_candidates)
                    if sr:
                        collect(sr, sn)
                        break
                except Exception:
                    pass
            if len(search_urls) >= max_candidates:
                break  # Have enough URLs to evaluate

        if not search_urls:
            # Try last-resort Google/Bing search before giving up
            send("status", {"message": f"🔍 No results from priority/general search, trying Google/Bing for '{denumire[:40]}'"})
            last_resort_queries_nr = [denumire, clean_product_query(denumire)]
            for lrq in dict.fromkeys(last_resort_queries_nr):
                for sn, searcher in [("bing", bing_scraper), ("ddg", ddg)]:
                    try:
                        lr = searcher.search(lrq, max_results=5)
                        if lr:
                            collect(lr, f"lastresort:{sn}")
                            break
                    except Exception:
                        pass
                if search_urls:
                    break
            if not search_urls:
                send("product_done", {
                    "index": idx, "product_id": product_id, "denumire": denumire,
                    "status": "no_results", "images": [], "source": "none",
                })
                stats["failed"] += 1
                results.append({"id": product_id, "denumire": denumire,
                                "status": "failed", "images": []})
                continue

        source_used = search_urls[0][1] if search_urls else "none"

        # ── DOWNLOAD ALL IN PARALLEL ──
        all_urls = [u for u, _ in search_urls]
        url_to_source = {u: s for u, s in search_urls}
        downloaded = download_images_parallel(all_urls, max_workers=6)

        send("status", {"message": f"Downloaded {len(downloaded)}/{len(all_urls)} images"})

        # ── EVALUATE: quality + CLIP, collect valid candidates ──
        # Track if priority sites returned anything — if not, relax general search filters
        has_priority_urls = any(s.startswith("direct:") for _, s in search_urls)
        saved_images = []
        candidates_tried = 0
        valid_candidates = []  # [(data, qc, relevance_score, url, img_hash, src)]
        seen_hashes = set()

        for url, src in search_urls:
            data = downloaded.get(url)
            if not data:
                continue
            candidates_tried += 1

            # Quality check — relaxed for priority site images
            is_priority = src.startswith("direct:")
            if is_priority:
                # Use a relaxed quality checker for priority site images
                # (many e-commerce sites have smaller product images)
                relaxed_checker = ImageQualityChecker({
                    "min_resolution": min(quality_checker.min_resolution, 500),
                    "min_quality_score": max(quality_checker.min_quality_score - 15, 25),
                    "reject_blurry": quality_checker.reject_blurry,
                    "min_aspect_ratio": quality_checker.min_aspect_ratio,
                    "max_aspect_ratio": quality_checker.max_aspect_ratio,
                })
                qc = relaxed_checker.evaluate(data)
            else:
                qc = quality_checker.evaluate(data)
            if not qc["passed"]:
                send("candidate_checked", {
                    "index": idx, "url": url[:100],
                    "quality_score": qc["score"], "passed": False,
                    "reasons": qc["reasons"], "details": qc["details"],
                    "relevance_score": None,
                })
                continue

            # Deduplicate
            img_hash = hashlib.md5(data).hexdigest()
            if img_hash in seen_hashes:
                continue
            seen_hashes.add(img_hash)

            # Relevance scoring: combine CLIP + URL keyword matching
            clip_score = 50
            # For priority site images, score against the PRODUCT PAGE URL (not the image URL)
            # Image URLs often use product codes (MON0116) or generic filenames,
            # while product page URLs contain the actual product slug
            score_url = url
            if is_priority and url in img_to_product_url:
                score_url = img_to_product_url[url]
            url_score = url_keyword_score(score_url, denumire)
            url_score_100 = round(url_score * 100)

            if relevance_checker:
                rc = relevance_checker.check(data, denumire)
                clip_score = rc["score"]

                # Note: CLIP multi_product flag is logged but NOT used for rejection.
                # Claude Vision API check (below) is far more reliable for this.
                if rc.get("multi_product", False):
                    logger.info(f"CLIP flagged multi-product (informational only): {url[:80]}")

                # Combined score: 70% CLIP + 30% URL keywords
                relevance_score = round(clip_score * 0.7 + url_score_100 * 0.3)

                if is_priority:
                    # Priority site images come from AI-matched product pages —
                    # their image URLs often use product codes (e.g. MON0116) instead
                    # of readable names, so url_keyword_score ≈ 0.  Don't penalize them.
                    # Never reject by CLIP alone; boost score so they win over general search.
                    relevance_score = max(relevance_score, clip_score, 75)

                    # But STILL check for conflicting products — even priority sites can
                    # return the wrong product (e.g. Green Jasmine instead of Blueberry Cream)
                    # Check the PRODUCT PAGE URL (not image URL which has infrastructure paths)
                    product_url_check = img_to_product_url.get(url, "")
                    has_conflict = False
                    if product_url_check:
                        has_conflict = url_has_conflicting_product(product_url_check, denumire)
                        if has_conflict:
                            print(f"[PRIORITY-FILTER] REJECT conflicting product_url: {product_url_check[:100]} for '{denumire}'")
                        # Also reject category pages: multi-segment paths like /cafea/cafea-boabe/lavazza
                        # are categories, not products. Products use single slugs like /lavazza-paulista-1kg
                        if not has_conflict and product_url_check:
                            purl_path = urlparse(product_url_check).path.strip("/")
                            purl_segments = [s for s in purl_path.split("/") if s]
                            if len(purl_segments) >= 2 and all(len(s) < 30 for s in purl_segments):
                                # Multi-segment, short segments = likely category page
                                # Check if product's distinctive words are in the URL
                                purl_text = purl_path.replace("/", " ").replace("-", " ").lower()
                                name_words_lc = set(re.findall(r'[a-z]{4,}', denumire.lower()))
                                _skip_generic = {"cafea", "coffee", "ceai", "tea", "boabe", "beans",
                                                 "capsule", "sirop", "syrup", "piure", "pireu"}
                                distinctive_chk = {w for w in name_words_lc
                                                   if w not in _skip_generic and w not in set(_RO_EN_MAP.keys()) | set(_RO_EN_MAP.values())}
                                if distinctive_chk:
                                    found_in_url = sum(1 for dw in distinctive_chk if dw in purl_text)
                                    if found_in_url < len(distinctive_chk):
                                        has_conflict = True
                                        print(f"[PRIORITY-FILTER] REJECT category page missing distinctive words: {product_url_check[:80]} (has {found_in_url}/{len(distinctive_chk)} of {distinctive_chk})")
                    if has_conflict:
                        send("candidate_checked", {
                            "index": idx, "url": url[:100],
                            "quality_score": qc["score"], "passed": False,
                            "reasons": ["Wrong product variant (priority site, URL conflict)"],
                            "details": {**qc["details"], "relevance": rc["similarity"],
                                        "url_match": url_score_100},
                            "relevance_score": relevance_score,
                        })
                        continue
                else:
                    # General search filtering:
                    # If priority sites had NO results, be lenient — accept the first
                    # image that passes quality + isn't multi-product. The Google/Bing
                    # query already contains the product name, so results are usually
                    # relevant enough.
                    if has_priority_urls:
                        # Priority sites had results, so general search should be strict
                        if not rc["relevant"] and url_score < 0.5:
                            send("candidate_checked", {
                                "index": idx, "url": url[:100],
                                "quality_score": qc["score"], "passed": False,
                                "reasons": [rc["reason"]],
                                "details": {**qc["details"], "relevance": rc["similarity"],
                                            "url_match": url_score_100},
                                "relevance_score": relevance_score,
                            })
                            continue
                    else:
                        # No priority site results — relaxed mode:
                        # Only reject if CLIP is VERY low AND URL has no keywords at all
                        if not rc["relevant"] and url_score < 0.2 and rc.get("similarity", 0) < 0.12:
                            send("candidate_checked", {
                                "index": idx, "url": url[:100],
                                "quality_score": qc["score"], "passed": False,
                                "reasons": [rc["reason"]],
                                "details": {**qc["details"], "relevance": rc["similarity"],
                                            "url_match": url_score_100},
                                "relevance_score": relevance_score,
                            })
                            continue
                        # Boost relevance score for general search when no priority found
                        # so these images don't get ranked below threshold
                        relevance_score = max(relevance_score, 50)

                    # CONFLICTING PRODUCT CHECK: reject if URL clearly indicates
                    # a DIFFERENT product variant (e.g. "jasmine" when searching "blueberry")
                    if url_has_conflicting_product(url, denumire):
                        send("candidate_checked", {
                            "index": idx, "url": url[:100],
                            "quality_score": qc["score"], "passed": False,
                            "reasons": ["Wrong product variant (URL has conflicting flavor/type)"],
                            "details": {**qc["details"], "relevance": rc["similarity"],
                                        "url_match": url_score_100},
                            "relevance_score": relevance_score,
                        })
                        continue
            else:
                relevance_score = url_score_100
                if is_priority:
                    # Without CLIP, priority site images still get a high baseline score
                    relevance_score = max(relevance_score, 75)
                elif not has_priority_urls:
                    # No priority results at all — boost general search baseline
                    relevance_score = max(relevance_score, 50)
                    # Check product_url for conflicts (not image URL)
                    product_url_check_nc = img_to_product_url.get(url, "")
                    if product_url_check_nc and url_has_conflicting_product(product_url_check_nc, denumire):
                        print(f"[PRIORITY-FILTER-NC] REJECT: {product_url_check_nc[:100]} for '{denumire}'")
                        send("candidate_checked", {
                            "index": idx, "url": url[:100],
                            "quality_score": qc["score"], "passed": False,
                            "reasons": ["Wrong product variant (priority site, URL conflict)"],
                            "details": {**qc["details"], "url_match": url_score_100},
                            "relevance_score": relevance_score,
                        })
                        continue
                elif not is_priority and url_has_conflicting_product(url, denumire):
                    send("candidate_checked", {
                        "index": idx, "url": url[:100],
                        "quality_score": qc["score"], "passed": False,
                        "reasons": ["Wrong product variant (URL has conflicting flavor/type)"],
                        "details": {**qc["details"], "url_match": url_score_100},
                        "relevance_score": relevance_score,
                    })
                    continue

            # ── GEMINI VISION: single product + text verification ──
            # Uses Google Gemini Flash (FREE) to verify:
            # 1. Only 1 product in image  2. Packaging text matches product name
            gemini_key = config.get("gemini_key", "") or os.environ.get("GEMINI_API_KEY", "") or _load_config_file().get("gemini_key", "")
            if gemini_key:
                send("status", {"message": f"👁️ Gemini Vision: {denumire[:40]}..."})
                vision_result = gemini_vision_check_image(data, denumire, gemini_key)
                v_ok = vision_result.get("ok", True)
                v_count = vision_result.get("count", "?")
                v_text = vision_result.get("visible_text", "")[:50]
                v_reason = vision_result.get("reason", "")
                if v_ok is True:
                    send("status", {"message": f"👁️ Gemini: {v_count} item, text='{v_text}' → OK"})
                elif v_ok is None:
                    # Gemini failed (429/error) — fallback to local checks
                    send("status", {"message": f"👁️ Gemini unavailable, using local checks..."})
                    # Use stricter CLIP + URL conflict as fallback
                    fallback_ok = True
                    fallback_reasons = []
                    # Check URL conflict (image URL + product page URL)
                    if url_has_conflicting_product(url, denumire):
                        fallback_ok = False
                        fallback_reasons.append("URL conflict (local fallback)")
                    product_url_fb = img_to_product_url.get(url, "")
                    if product_url_fb and url_has_conflicting_product(product_url_fb, denumire):
                        fallback_ok = False
                        fallback_reasons.append("Product page conflict (local fallback)")
                    # Require higher CLIP score when no Vision
                    if clip_score is not None and clip_score < 0.22:
                        fallback_ok = False
                        fallback_reasons.append(f"Low CLIP score {clip_score:.2f} (local fallback)")
                    if not fallback_ok:
                        send("candidate_checked", {
                            "index": idx, "url": url[:100],
                            "quality_score": qc["score"], "passed": False,
                            "reasons": fallback_reasons,
                            "details": {**qc["details"], "url_match": url_score_100, "clip": clip_score},
                            "relevance_score": relevance_score,
                        })
                        continue
                    send("status", {"message": f"👁️ Local fallback: CLIP={clip_score}, URL OK → accepted"})
                else:
                    send("status", {"message": f"👁️ Gemini REJECT: {v_reason}"})
                    send("candidate_checked", {
                        "index": idx, "url": url[:100],
                        "quality_score": qc["score"], "passed": False,
                        "reasons": [f"Gemini Vision: {v_reason}"],
                        "details": {**qc["details"], "url_match": url_score_100,
                                    "clip": clip_score,
                                    "vision_count": v_count,
                                    "vision_text": v_text},
                        "relevance_score": relevance_score,
                    })
                    continue

            send("candidate_checked", {
                "index": idx, "url": url[:100],
                "quality_score": qc["score"], "passed": True,
                "reasons": [],
                "details": {**qc["details"], "url_match": url_score_100,
                            "clip": clip_score},
                "relevance_score": relevance_score,
            })

            valid_candidates.append((data, qc, relevance_score, url, img_hash, src))

            # EARLY STOP: only if we have enough great matches for ALL requested images
            if len(valid_candidates) >= images_per_product:
                great_count = sum(1 for c in valid_candidates if c[2] >= GREAT_RELEVANCE)
                if great_count >= images_per_product:
                    break
                # Without CLIP: stop after enough candidates
                if not relevance_checker:
                    break

        # Sort by relevance (highest first), with priority site images winning ties
        # Also preserve matcher ordering: earlier results from priority sites get a position bonus
        # (the matcher already ranked them by relevance — first result = best match)
        def _sort_key(c):
            is_pri = 1 if c[5].startswith("direct:") else 0
            rel = c[2]  # relevance_score
            # Position bonus: first priority result gets +10, second +5, etc.
            # This preserves the matcher's ranking when URL keyword scores are close
            url = c[3]
            try:
                pos_in_search = next(i for i, (u, _) in enumerate(search_urls) if u == url)
                if is_pri:
                    position_bonus = max(0, 15 - pos_in_search * 5)
                else:
                    position_bonus = 0
            except StopIteration:
                position_bonus = 0
            return (is_pri, rel + position_bonus)

        valid_candidates.sort(key=_sort_key, reverse=True)

        # ── LAST RESORT: if no valid candidates, grab first decent Google/Bing image ──
        if not valid_candidates:
            send("status", {"message": f"🔍 Last resort: searching Google/Bing for '{denumire[:40]}'"})
            last_resort_queries = [denumire, clean_product_query(denumire)]
            for lrq in dict.fromkeys(last_resort_queries):
                if valid_candidates:
                    break
                for sn, searcher in [("bing", bing_scraper), ("ddg", ddg)]:
                    try:
                        lr_results = searcher.search(lrq, max_results=5)
                        if not lr_results:
                            continue
                        lr_urls = [r.get("image", "") for r in lr_results if r.get("image")]
                        lr_downloaded = download_images_parallel(lr_urls, max_workers=4)
                        for lr_url in lr_urls:
                            lr_data = lr_downloaded.get(lr_url)
                            if not lr_data:
                                continue
                            lr_qc = quality_checker.evaluate(lr_data)
                            if not lr_qc["passed"]:
                                continue
                            lr_hash = hashlib.md5(lr_data).hexdigest()
                            if lr_hash in seen_hashes:
                                continue
                            # Accept first image that passes quality — skip Gemini/CLIP for last resort
                            print(f"[LAST-RESORT] Accepting first decent image from {sn}: {lr_url[:100]}")
                            send("status", {"message": f"🔍 Last resort: found image from {sn}"})
                            valid_candidates.append((lr_data, lr_qc, 50, lr_url, lr_hash, f"lastresort:{sn}"))
                            break
                        if valid_candidates:
                            break
                    except Exception as e:
                        logger.debug(f"Last resort search error ({sn}): {e}")

        for data, qc, rel_score, url, img_hash, src in valid_candidates[:images_per_product]:
            try:
                processed = resize_and_pad(data, target_size, quality=quality,
                                           remove_bg=remove_bg,
                                           output_format=output_format)
                img_num = len(saved_images) + 1
                filename = hermes_filename(product_id, denumire, img_num,
                                           fmt=output_format)
                filepath = output_dir / filename
                filepath.write_bytes(processed)

                thumb = make_thumbnail(processed)

                try:
                    image_domain = urlparse(url).netloc.replace("www.", "")
                except Exception:
                    image_domain = ""

                saved_images.append({
                    "filename": filename,
                    "quality_score": qc["score"],
                    "relevance_score": rel_score,
                    "source": src,
                    "image_url": url,
                    "image_domain": image_domain,
                    "hash": img_hash,
                    "thumbnail": thumb,
                    "details": qc["details"],
                })
            except Exception as e:
                logger.debug(f"Process error for {denumire}: {e}")

        status = "ok" if saved_images else "failed"
        if saved_images:
            stats["success"] += 1
            stats["images_saved"] += len(saved_images)
        else:
            stats["failed"] += 1

        # Use the winning candidate's source, not the first URL's source
        actual_source = saved_images[0]["source"] if saved_images else source_used
        send("product_done", {
            "index": idx, "product_id": product_id, "denumire": denumire,
            "status": status, "images": saved_images, "source": actual_source,
            "candidates_tried": candidates_tried,
        })

        results.append({
            "id": product_id, "denumire": denumire,
            "status": status, "images": saved_images,
        })

        # Minimal delay to avoid rate limiting
        if idx < len(products) - 1:
            time.sleep(0.1)

    send("job_done", {
        "job_id": job_id,
        "stats": stats,
        "output_dir": str(output_dir.resolve()),
    })


# ─── FILE PARSERS ─────────────────────────────────────────────────────────

def parse_uploaded_file(file_storage) -> list[str]:
    """
    Extract product names from uploaded file.
    Supports: .xlsx, .xls, .csv, .tsv, .txt, .docx, .pdf
    Returns a list of product name strings.
    """
    filename = file_storage.filename.lower()
    data = file_storage.read()

    if filename.endswith((".xlsx", ".xls")):
        return _parse_excel(data)
    elif filename.endswith(".csv"):
        return _parse_csv(data, ",")
    elif filename.endswith(".tsv"):
        return _parse_csv(data, "\t")
    elif filename.endswith(".txt"):
        return _parse_txt(data)
    elif filename.endswith(".docx"):
        return _parse_docx(data)
    elif filename.endswith(".pdf"):
        return _parse_pdf(data)
    else:
        raise ValueError(f"Unsupported file type: {filename}")


def _parse_excel(data: bytes) -> list[str]:
    """Parse .xlsx - takes the first column with text data, or a 'denumire' column."""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    # Check header row for a known column name
    header = [str(c).lower().strip() if c else "" for c in rows[0]]
    target_col = None
    for keyword in ["denumire", "produs", "product", "name", "nume", "descriere", "description", "item"]:
        if keyword in header:
            target_col = header.index(keyword)
            break

    if target_col is not None:
        # Use the identified column, skip header
        lines = [str(row[target_col]).strip() for row in rows[1:] if row[target_col]]
    else:
        # Auto-detect: use the first column that has mostly text
        # Try each column, pick the one with most non-empty string values
        best_col = 0
        best_count = 0
        num_cols = max(len(r) for r in rows) if rows else 0
        for ci in range(num_cols):
            count = sum(1 for r in rows if len(r) > ci and r[ci]
                        and isinstance(r[ci], str) and len(str(r[ci]).strip()) > 2)
            if count > best_count:
                best_count = count
                best_col = ci

        # Check if first row looks like a header
        first_val = str(rows[0][best_col]).strip() if rows[0][best_col] else ""
        start = 1 if (first_val.lower() in header or len(first_val) < 20) else 0
        lines = [str(row[best_col]).strip() for row in rows[start:]
                 if len(row) > best_col and row[best_col] and str(row[best_col]).strip()]

    wb.close()
    return [l for l in lines if l and l.lower() != "none"]


def _parse_csv(data: bytes, delimiter: str) -> list[str]:
    """Parse CSV/TSV - same logic as Excel: find 'denumire' column or first text column."""
    import csv as csv_mod
    text = data.decode("utf-8-sig", errors="replace")
    # Auto-detect delimiter if comma doesn't produce multiple columns
    reader = list(csv_mod.reader(io.StringIO(text), delimiter=delimiter))
    if not reader:
        return []

    # Also try semicolon (common in RO)
    if delimiter == "," and all(len(r) <= 1 for r in reader[:5]):
        reader = list(csv_mod.reader(io.StringIO(text), delimiter=";"))

    header = [str(c).lower().strip() for c in reader[0]] if reader else []
    target_col = None
    for keyword in ["denumire", "produs", "product", "name", "nume", "descriere", "description", "item"]:
        if keyword in header:
            target_col = header.index(keyword)
            break

    if target_col is not None:
        return [row[target_col].strip() for row in reader[1:] if len(row) > target_col and row[target_col].strip()]

    # Fallback: first column with longest average text
    if reader:
        num_cols = max(len(r) for r in reader)
        best_col = 0
        best_avg = 0
        for ci in range(num_cols):
            vals = [r[ci] for r in reader[1:] if len(r) > ci and r[ci].strip()]
            avg = sum(len(v) for v in vals) / max(len(vals), 1)
            if avg > best_avg:
                best_avg = avg
                best_col = ci
        return [row[best_col].strip() for row in reader[1:] if len(row) > best_col and row[best_col].strip()]

    return []


def _parse_txt(data: bytes) -> list[str]:
    """Parse plain text - one product per line."""
    text = data.decode("utf-8-sig", errors="replace")
    return [line.strip() for line in text.splitlines() if line.strip()]


def _parse_docx(data: bytes) -> list[str]:
    """Parse .docx - extract paragraphs and table cells."""
    from docx import Document
    doc = Document(io.BytesIO(data))
    lines = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and len(text) > 2:
            lines.append(text)

    # Tables (often product lists are in tables)
    for table in doc.tables:
        header = [cell.text.lower().strip() for cell in table.rows[0].cells] if table.rows else []
        target_col = None
        for keyword in ["denumire", "produs", "product", "name", "nume"]:
            if keyword in header:
                target_col = header.index(keyword)
                break

        start_row = 0
        if target_col is not None:
            start_row = 1
        else:
            # Use first column with longest text
            target_col = 0

        for row in table.rows[start_row:]:
            cells = row.cells
            if len(cells) > target_col:
                text = cells[target_col].text.strip()
                if text and len(text) > 2:
                    lines.append(text)

    return lines


def _parse_pdf(data: bytes) -> list[str]:
    """Parse .pdf - extract text lines."""
    import fitz  # pymupdf
    doc = fitz.open(stream=data, filetype="pdf")
    lines = []
    for page in doc:
        text = page.get_text()
        for line in text.splitlines():
            line = line.strip()
            if line and len(line) > 2:
                lines.append(line)
    doc.close()
    return lines


# ─── API ROUTES ───────────────────────────────────────────────────────────

CONFIG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")

def _load_config_file() -> dict:
    """Load saved config from config.json."""
    try:
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as f:
                return json.load(f)
    except Exception as e:
        logger.warning(f"Failed to load config.json: {e}")
    return {}

def _save_config_file(data: dict):
    """Save config to config.json."""
    try:
        existing = _load_config_file()
        existing.update(data)
        with open(CONFIG_FILE, "w") as f:
            json.dump(existing, f, indent=2)
    except Exception as e:
        logger.warning(f"Failed to save config.json: {e}")


@app.route("/")
def index():
    return send_from_directory(".", "index.html")


@app.route("/api/config", methods=["GET"])
def get_config():
    """Return saved config (with keys partially masked)."""
    cfg = _load_config_file()
    # Return masked keys for display
    result = {}
    for k, v in cfg.items():
        if "key" in k.lower() and v:
            result[k] = v  # Frontend will handle masking
        else:
            result[k] = v
    return jsonify(result)


@app.route("/api/config", methods=["POST"])
def save_config():
    """Save config values to config.json."""
    data = request.json or {}
    # Only save non-empty keys
    to_save = {k: v for k, v in data.items() if v}
    _save_config_file(to_save)
    return jsonify({"ok": True})


@app.route("/api/upload", methods=["POST"])
def upload_file():
    """Upload a file and extract product names from it."""
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "Empty filename"}), 400

    try:
        products = parse_uploaded_file(f)
        return jsonify({
            "products": products,
            "count": len(products),
            "filename": f.filename,
        })
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        logger.error(f"File parse error: {e}")
        return jsonify({"error": f"Failed to parse file: {str(e)}"}), 500


@app.route("/api/stop/<job_id>", methods=["POST"])
def stop_job(job_id):
    """Cancel a running job."""
    job = active_jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    job["status"] = "cancelled"
    return jsonify({"ok": True, "job_id": job_id})


@app.route("/api/start", methods=["POST"])
def start_job():
    """Start a scraping job. Expects JSON with products list and config."""
    payload = request.json
    products = payload.get("products", [])
    config = payload.get("config", {})

    if not products:
        return jsonify({"error": "No products provided"}), 400

    job_id = datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:6]
    event_queue = queue.Queue()

    thread = threading.Thread(
        target=run_scraper_job,
        args=(job_id, products, config, event_queue),
        daemon=True,
    )

    active_jobs[job_id] = {
        "thread": thread,
        "queue": event_queue,
        "status": "running",
        "config": config,
    }

    thread.start()
    return jsonify({"job_id": job_id})


@app.route("/api/stream/<job_id>")
def stream_events(job_id):
    """SSE endpoint for real-time progress updates."""
    job = active_jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404

    def generate():
        while True:
            try:
                msg = job["queue"].get(timeout=30)
                yield f"data: {json.dumps(msg)}\n\n"
                if msg.get("event") == "job_done":
                    job["status"] = "done"
                    break
            except queue.Empty:
                yield "data: {\"event\": \"heartbeat\"}\n\n"
                if not job["thread"].is_alive():
                    yield f'data: {json.dumps({"event": "job_done", "data": {"error": "Thread died"}})}\n\n'
                    break

    return Response(generate(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/api/images/<job_id>/<filename>")
def serve_image(job_id, filename):
    """Serve a downloaded image."""
    img_dir = OUTPUT_DIR / job_id
    return send_from_directory(str(img_dir), filename)


@app.route("/api/jobs")
def list_jobs():
    """List all jobs and their status."""
    return jsonify({
        jid: {"status": j["status"], "config": j["config"]}
        for jid, j in active_jobs.items()
    })


# ─── MAIN ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Image Scraper UI")
    parser.add_argument("--port", "-p", type=int, default=8787)
    parser.add_argument("--host", default="0.0.0.0")
    parser.add_argument("--debug", action="store_true")
    args = parser.parse_args()

    print(f"\n  Image Scraper UI running at http://localhost:{args.port}")
    print(f"  [v3.7] + context-aware cream + product line check on product_url + rescue search\n")
    app.run(host=args.host, port=args.port, debug=args.debug, threaded=True)
